from pathlib import Path
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


DOCS_DIR = Path(r"c:\Users\Junior Owusu\Desktop\Final Year Project\child-vaccination-system\docs")
DIAGRAMS_DIR = Path(r"c:\Users\Junior Owusu\Desktop\Final Year Project\diagrams")

SOURCE_MD = DOCS_DIR / "CHAPTER-3-DRAFT.md"
OUT_DOCX = DOCS_DIR / "CHAPTER-3.docx"


HEADING_CITATIONS = {
    "3.0": [2, 4],
    "3.1": [30, 1, 2],
    "3.2": [9, 10],
    "3.3": [11, 13, 14],
    "3.4": [11, 13, 14],
    "3.5": [13, 14],
    "3.6": [13, 29],
    "3.7": [2, 9, 13],
    "3.8": [9, 10],
    "3.9": [11, 13, 14],
}


TEXT_CITATION_RULES: list[tuple[re.Pattern[str], list[int]]] = [
    (re.compile(r"html5-qrcode|Html5Qrcode|QrScanner|rear-facing camera|camera lifecycle", re.IGNORECASE), [23]),
    (re.compile(r"qr_code_payload|signed token|cvcc_id|CH-\d{4}-\d{3}|Quick Response \(QR\).*payload", re.IGNORECASE), [20, 23, 25]),
    (re.compile(r"transfer-out|transfer-in|TransferOutDto|TransferInDto|primary_facility_id|queueTransferIn|queueTransferOut", re.IGNORECASE), [13, 14, 15]),
    (re.compile(r"dispatches notifications.*SMS.*email|notification system", re.IGNORECASE), [21, 22]),
    (re.compile(r"Hubtel|SmsService|welcome SMS|vaccination reminder", re.IGNORECASE), [21]),
    (re.compile(r"Brevo|transactional email", re.IGNORECASE), [22]),
    (re.compile(r"jsPDF|generateCertificatePdf|CERT-GH|tamper-evident PDF|vaccination certificate", re.IGNORECASE), [24, 25, 23]),
    (re.compile(r"Leaflet|React-Leaflet|OpenStreetMap|tile source|map tiles|outreach-map", re.IGNORECASE), [26, 27, 28]),
    (re.compile(r"Dexie|IndexedDB|offline-first|service worker|offline conditions|network restoration|idempotencyKey", re.IGNORECASE), [16, 17, 18]),
    (re.compile(r"Next\.js|React|App Router", re.IGNORECASE), [11, 12]),
    (re.compile(r"NestJS|Data Transfer Object|DTO|class-validator|controller", re.IGNORECASE), [13, 29]),
    (re.compile(r"Supabase|PostgreSQL|PostGIS|POLYGON|POINT values|GPS coordinates", re.IGNORECASE), [14, 15, 28]),
    (re.compile(r"Role-Based Access Control|RBAC|Json Web Token|JWT|Row-Level Security|RLS", re.IGNORECASE), [13, 20]),
    (re.compile(r"semi-structured|phone calls|field research|nurses and community health workers", re.IGNORECASE), [30, 1, 2]),
    (re.compile(r"Agile|Sprint|Scrum|retrospective|product backlog", re.IGNORECASE), [9, 10]),
    (re.compile(r"immunization|public health|cross-facility", re.IGNORECASE), [2, 4]),
]


PREFIX_CITATION_RULES: list[tuple[str, list[int]]] = [
    ("The chapter begins by describing how data was collected from nurses", [30, 1, 2]),
    ("Through these interviews, several recurring pain points were identified", [30, 1, 2]),
    ("The second constraint was catchment area rules", [4, 30]),
    ("The project was built by a team of three developers", [9, 10]),
    ("The CVCC system is built on a three-tier architecture", [11, 13, 14, 15]),
    ("The flow of a typical interaction", [13, 20, 29]),
    ("Synchronisation is handled by the ChwBackgroundSyncService", [16, 17, 18]),
    ("Geographical mapping for CHW outreach scheduling is powered by Leaflet", [26, 27, 28]),
    ("The CVCC system manages seven distinct user roles", [13, 20]),
    ("At the database layer, Supabase Row-Level Security", [14, 15, 20]),
    ("The persistent data store for the CVCC system is a PostgreSQL database", [14, 15, 28]),
    ("Background job scheduling does not use an external queue server", [13, 21]),
    ("On the scanning side, the Html5Qrcode component in the PHA portal", [23, 13, 14]),
    ("The embedded QR code is generated from the certificates.qr_payload value", [24, 25, 23]),
    ("The quality assurance strategy for the CVCC system", [13, 29]),
    ("With the device simulated as offline", [16, 17, 18]),
    ("Bugs identified during the development process were tracked", [9, 10]),
    ("This chapter has documented the methodology and design decisions", [2, 9, 13]),
    ("Sprint execution:", [9, 10]),
]


EXPANSION_PATTERNS = [
    (re.compile(r"(?<!\()RBAC\b"), "Role-Based Access Control (RBAC)", "RBAC"),
    (re.compile(r"(?<!\()PWA\b"), "Progressive Web Application (PWA)", "PWA"),
    (re.compile(r"(?<!\()RLS\b"), "Row-Level Security (RLS)", "RLS"),
    (re.compile(r"\bSDLC\b"), "Software Development Life Cycle (SDLC)", "SDLC"),
    (re.compile(r"\bRESTful API\b"), "Representational State Transfer (REST) Application Programming Interface (API)", "API"),
    (re.compile(r"\bAPI\b"), "Application Programming Interface (API)", "API"),
    (re.compile(r"\bJWT token\b"), "Json Web Token (JWT)", "JWT"),
    (re.compile(r"\bJWT\b"), "Json Web Token (JWT)", "JWT"),
    (re.compile(r"\bSMS\b"), "Short Message Service (SMS)", "SMS"),
    (re.compile(r"\bPDF\b"), "Portable Document Format (PDF)", "PDF"),
    (re.compile(r"\bQR code\b"), "Quick Response (QR) code", "QR"),
    (re.compile(r"\bQR\b"), "Quick Response (QR)", "QR"),
    (re.compile(r"\bDTOs\b"), "Data Transfer Objects (DTOs)", "DTO"),
]


expanded_flags: set[str] = set()


def configure_document(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)

    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.0)


def clean_inline(text: str) -> str:
    # Remove markdown emphasis and inline code markers for Word rendering.
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    text = re.sub(r"`([^`]*)`", r"\1", text)
    text = text.replace("\u2014", "-")
    text = text.replace("\u2013", "-")
    text = text.replace("\u201c", '"').replace("\u201d", '"')
    text = text.replace("\u2019", "'")
    return text.strip()


def get_heading_key(raw_heading: str) -> str | None:
    text = raw_heading.strip()
    m = re.match(r"^(\d+\.\d+(?:\.\d+)?)", text)
    if not m:
        return None
    return m.group(1)


def expand_acronyms_once(text: str) -> str:
    updated = text
    for pattern, replacement, key in EXPANSION_PATTERNS:
        if key in expanded_flags:
            continue
        if pattern.search(updated):
            updated = pattern.sub(replacement, updated, count=1)
            expanded_flags.add(key)
    return updated


def append_ieee_citation(text: str, heading_key: str | None) -> str:
    cleaned = text.strip()
    if re.search(r"\[\d+\]", cleaned):
        return cleaned

    if heading_key and heading_key.startswith("3.8"):
        refs_text = ", ".join(f"[{num}]" for num in HEADING_CITATIONS["3.8"])
        if cleaned.endswith("."):
            cleaned = cleaned[:-1].rstrip()
        return f"{cleaned} {refs_text}."

    for prefix, mapped_refs in PREFIX_CITATION_RULES:
        if cleaned.startswith(prefix):
            refs_text = ", ".join(f"[{num}]" for num in mapped_refs)
            if cleaned.endswith("."):
                cleaned = cleaned[:-1].rstrip()
            return f"{cleaned} {refs_text}."

    refs: list[int] | None = None

    for pattern, mapped_refs in TEXT_CITATION_RULES:
        if pattern.search(cleaned):
            refs = mapped_refs
            break

    if refs is None and heading_key:
        refs = HEADING_CITATIONS.get(heading_key)
        if refs is None and "." in heading_key:
            section_key = ".".join(heading_key.split(".")[:2])
            refs = HEADING_CITATIONS.get(section_key)

    if refs is None:
        return cleaned

    refs_text = ", ".join(f"[{num}]" for num in refs)

    if cleaned.endswith("."):
        cleaned = cleaned[:-1].rstrip()

    return f"{cleaned} {refs_text}."


def add_chapter_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(16)


def add_heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = True

    if level == 2:
        run.font.size = Pt(14)
    elif level == 3:
        run.font.size = Pt(13)
    else:
        run.font.size = Pt(12)


def add_body(doc: Document, text: str, heading_key: str | None = None) -> None:
    content = expand_acronyms_once(text)
    content = append_ieee_citation(content, heading_key)

    p = doc.add_paragraph(content)
    p.paragraph_format.first_line_indent = Inches(0.3)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(6)


def add_bullet(doc: Document, text: str, heading_key: str | None = None) -> None:
    content = expand_acronyms_once(text)
    content = append_ieee_citation(content, heading_key)

    p = doc.add_paragraph(content)
    p.style = doc.styles["List Bullet"]
    p.paragraph_format.line_spacing = 1.5


def add_mermaid_code_block(doc: Document, lines: list[str]) -> None:
    for line in lines:
        p = doc.add_paragraph(line)
        p.paragraph_format.left_indent = Inches(0.4)
        p.paragraph_format.space_after = Pt(0)
        run = p.runs[0]
        run.font.name = "Courier New"
        run.font.size = Pt(10)


def add_system_architecture_diagram(doc: Document) -> None:
    diagram_path = DIAGRAMS_DIR / "system architectural diagram.png"
    if not diagram_path.exists():
        return

    add_body(
        doc,
        "System Architectural Diagram for the Child Vaccination Command Center (CVCC).",
        heading_key="3.3.1",
    )
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(diagram_path), width=Inches(6.0))


def parse_markdown_to_doc(doc: Document, markdown: str) -> None:
    lines = markdown.splitlines()

    buffer: list[str] = []
    in_code_block = False
    current_heading_key: str | None = None
    architecture_inserted = False

    def flush_buffer() -> None:
        nonlocal buffer
        if buffer:
            paragraph = clean_inline(" ".join(buffer).strip())
            if paragraph:
                add_body(doc, paragraph, heading_key=current_heading_key)
            buffer = []

    for raw in lines:
        line = raw.rstrip()
        stripped = line.strip()

        if stripped.startswith("```"):
            flush_buffer()
            in_code_block = not in_code_block
            continue

        if in_code_block:
            p = doc.add_paragraph(stripped)
            p.paragraph_format.left_indent = Inches(0.4)
            p.paragraph_format.space_after = Pt(0)
            run = p.runs[0]
            run.font.name = "Courier New"
            run.font.size = Pt(10)
            continue

        if not stripped:
            flush_buffer()
            continue

        if stripped == "---":
            flush_buffer()
            continue

        if stripped.startswith("# "):
            flush_buffer()
            add_chapter_title(doc, clean_inline(stripped[2:]))
            continue

        if stripped.startswith("## "):
            flush_buffer()
            heading = clean_inline(stripped[3:])
            add_heading(doc, heading, level=2)
            heading_key = get_heading_key(heading)
            if heading_key:
                current_heading_key = heading_key
            continue

        if stripped.startswith("### "):
            flush_buffer()
            heading = clean_inline(stripped[4:])
            add_heading(doc, heading, level=3)
            heading_key = get_heading_key(heading)
            if heading_key:
                current_heading_key = heading_key
            if heading.startswith("3.3.1") and not architecture_inserted:
                add_system_architecture_diagram(doc)
                architecture_inserted = True
            continue

        if stripped.startswith("#### "):
            flush_buffer()
            heading = clean_inline(stripped[5:])
            add_heading(doc, heading, level=4)
            heading_key = get_heading_key(heading)
            if heading_key:
                current_heading_key = heading_key
            continue

        if stripped.startswith("- "):
            flush_buffer()
            add_bullet(doc, clean_inline(stripped[2:]), heading_key=current_heading_key)
            continue

        buffer.append(stripped)

    flush_buffer()


def append_sprint_methodology(doc: Document) -> None:
    add_heading(doc, "3.8 Sprint Methodology Integration", level=2)
    add_body(
        doc,
        "To strengthen the Agile approach used in this project, the team adopted a Sprint workflow that organized implementation into short, iterative cycles. Each Sprint focused on a prioritized group of user stories, with planning, implementation, review, and retrospective activities completed before the next cycle began.",
        heading_key="3.8",
    )
    add_bullet(
        doc,
        "Sprint planning: user stories, acceptance criteria, and dependency identification.",
        heading_key="3.8",
    )
    add_bullet(
        doc,
        "Sprint execution: frontend and backend implementation with daily synchronization.",
        heading_key="3.8",
    )
    add_bullet(
        doc,
        "Sprint review: demonstration of completed functionality against agreed outcomes.",
        heading_key="3.8",
    )
    add_bullet(
        doc,
        "Sprint retrospective: defect analysis and process improvements for the next Sprint.",
        heading_key="3.8",
    )

    add_heading(doc, "3.8.1 Sprint Flow in Mermaid Code", level=3)
    add_body(
        doc,
        "The Sprint cycle used by the project team is represented below in Mermaid format.",
        heading_key="3.8",
    )
    add_mermaid_code_block(
        doc,
        [
            "flowchart LR",
            "    A[Product Backlog - all pending tasks] --> B[Sprint Planning - choose tasks for this cycle]",
            "    B --> C[Sprint Backlog - selected tasks for this Sprint]",
            "    C --> D[Sprint Execution - build and test the selected tasks]",
            "    D --> E[Daily Standup - quick team progress check]",
            "    E --> D",
            "    D --> F[Sprint Review - demonstrate completed work]",
            "    F --> G[Sprint Retrospective - discuss improvements for next Sprint]",
            "    G --> H[Increment Release - deliver working features]",
            "    H --> A",
        ],
    )


def append_all_diagrams(doc: Document) -> None:
    add_heading(doc, "3.9 System Diagrams", level=2)
    add_body(
        doc,
        "This section includes all project system diagrams from the diagrams folder as requested.",
        heading_key="3.9",
    )

    diagram_files = sorted(DIAGRAMS_DIR.glob("*.png"))
    if not diagram_files:
        add_body(
            doc,
            "No diagrams were found in the configured diagrams directory.",
            heading_key="3.9",
        )
        return

    for idx, img_path in enumerate(diagram_files, start=1):
        caption = img_path.stem.replace("-", " ").replace("_", " ").strip().title()
        add_heading(doc, f"3.9.{idx} {caption}", level=3)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(img_path), width=Inches(6.0))


def append_reference_note(doc: Document) -> None:
    p = doc.add_paragraph(
        "Note: All references cited in this chapter are consolidated in CHAPTER FIVE REFERENCES."
    )
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(0)
    if p.runs:
        p.runs[0].italic = True


def main() -> None:
    if not SOURCE_MD.exists():
        raise FileNotFoundError(f"Source file not found: {SOURCE_MD}")

    markdown_text = SOURCE_MD.read_text(encoding="utf-8", errors="ignore")

    doc = Document()
    configure_document(doc)

    parse_markdown_to_doc(doc, markdown_text)
    append_sprint_methodology(doc)
    append_all_diagrams(doc)
    append_reference_note(doc)

    try:
        doc.save(OUT_DOCX)
        print(f"Regenerated detailed chapter: {OUT_DOCX}")
    except PermissionError:
        alt_path = DOCS_DIR / "CHAPTER-3-UPDATED.docx"
        doc.save(alt_path)
        print(f"Primary output was locked. Saved updated chapter to: {alt_path}")


if __name__ == "__main__":
    main()
