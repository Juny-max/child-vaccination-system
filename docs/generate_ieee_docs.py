from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


WORKSPACE_DOCS = Path(r"c:\Users\Junior Owusu\Desktop\Final Year Project\child-vaccination-system\docs")
DIAGRAMS_DIR = Path(r"c:\Users\Junior Owusu\Desktop\Final Year Project\diagrams")

CHAPTER1_PATH = WORKSPACE_DOCS / "CHAPTER ONE  (1).docx"
CHAPTER2_PATH = WORKSPACE_DOCS / "CHAPTER 2.docx"
CHAPTER3_PATH = WORKSPACE_DOCS / "CHAPTER-3.docx"
REFERENCES_PATH = WORKSPACE_DOCS / "CHAPTER 5 REFERENCES.docx"


REFERENCES = [
    "World Health Organization, \"Immunization coverage,\" 2025. [Online]. Available: https://www.who.int/news-room/fact-sheets/detail/immunization-coverage. [Accessed: Apr. 10, 2026].",
    "World Health Organization, \"Immunization Agenda 2030: A global strategy to leave no one behind,\" 2020. [Online]. Available: https://www.who.int/initiatives/immunization-agenda-2030. [Accessed: Apr. 10, 2026].",
    "United Nations Children's Fund, \"The State of the World's Children: For every child, vaccination,\" 2023. [Online]. Available: https://www.unicef.org/reports/state-worlds-children. [Accessed: Apr. 10, 2026].",
    "Ghana Health Service, \"Expanded Programme on Immunization operational guidance,\" Accra, Ghana, 2024.",
    "World Health Organization, \"Global strategy on digital health 2020-2025,\" 2021. [Online]. Available: https://www.who.int/publications/i/item/9789240020924. [Accessed: Apr. 28, 2026].",
    "World Health Organization, \"Recommendations on digital interventions for health system strengthening,\" WHO Guideline, 2019. [Online]. Available: https://www.who.int/publications/i/item/9789241550505. [Accessed: Apr. 28, 2026].",
    "Centers for Disease Control and Prevention, \"Immunization Information Systems (IIS),\" 2025. [Online]. Available: https://www.cdc.gov/vaccines/programs/iis/index.html. [Accessed: Apr. 28, 2026].",
    "Centers for Disease Control and Prevention, \"Immunization Information Systems (IIS) Technical Guidance,\" 2025. [Online]. Available: https://www.cdc.gov/vaccines/programs/iis/technical-guidance/index.html. [Accessed: Apr. 28, 2026].",
    "DHIS2, \"Individual Data Records with Tracker & Event Programs,\" 2025. [Online]. Available: https://dhis2.org/tracker. [Accessed: Apr. 28, 2026].",
    "F. D. Davis, \"Perceived usefulness, perceived ease of use, and user acceptance of information technology,\" MIS Quarterly, vol. 13, no. 3, pp. 319-340, 1989.",
    "V. Venkatesh, M. G. Morris, G. B. Davis, and F. D. Davis, \"User acceptance of information technology: Toward a unified view,\" MIS Quarterly, vol. 27, no. 3, pp. 425-478, 2003.",
    "A. Glasgow, T. Vogt, and S. Boles, \"Evaluating the public health impact of health promotion interventions: The RE-AIM framework,\" American Journal of Public Health, vol. 89, no. 9, pp. 1322-1327, 1999.",
    "DHIS2, \"Tracker and Event Program documentation,\" 2025. [Online]. Available: https://docs.dhis2.org. [Accessed: Apr. 10, 2026].",
    "K. Beck et al., \"Manifesto for Agile Software Development,\" 2001. [Online]. Available: https://agilemanifesto.org. [Accessed: Apr. 10, 2026].",
    "K. Schwaber and J. Sutherland, \"The Scrum Guide,\" 2020. [Online]. Available: https://scrumguides.org. [Accessed: Apr. 10, 2026].",
    "Vercel, \"Next.js documentation,\" 2025. [Online]. Available: https://nextjs.org/docs. [Accessed: Apr. 10, 2026].",
    "Meta Open Source, \"React documentation,\" 2025. [Online]. Available: https://react.dev. [Accessed: Apr. 10, 2026].",
    "NestJS Core Team, \"NestJS documentation,\" 2025. [Online]. Available: https://docs.nestjs.com. [Accessed: Apr. 10, 2026].",
    "Supabase Inc., \"Supabase documentation,\" 2025. [Online]. Available: https://supabase.com/docs. [Accessed: Apr. 10, 2026].",
    "PostgreSQL Global Development Group, \"PostgreSQL documentation,\" 2025. [Online]. Available: https://www.postgresql.org/docs. [Accessed: Apr. 10, 2026].",
    "D. Fahlander, \"Dexie.js documentation,\" 2025. [Online]. Available: https://dexie.org/docs. [Accessed: Apr. 10, 2026].",
    "W3C, \"Indexed Database API 3.0,\" 2022. [Online]. Available: https://www.w3.org/TR/IndexedDB. [Accessed: Apr. 10, 2026].",
    "Google, \"Progressive Web Applications,\" 2025. [Online]. Available: https://web.dev/progressive-web-apps. [Accessed: Apr. 10, 2026].",
    "National Institute of Standards and Technology, \"Advanced Encryption Standard (AES),\" FIPS PUB 197, 2001.",
    "M. Jones, J. Bradley, and N. Sakimura, \"Json Web Token (JWT),\" IETF RFC 7519, 2015.",
    "Hubtel, \"Hubtel SMS API documentation,\" 2025. [Online]. Available: https://developers.hubtel.com. [Accessed: Apr. 10, 2026].",
    "Brevo, \"Brevo Transactional Email API documentation,\" 2025. [Online]. Available: https://developers.brevo.com. [Accessed: Apr. 10, 2026].",
    "M. Dhar, \"html5-qrcode,\" 2025. [Online]. Available: https://github.com/mebjas/html5-qrcode. [Accessed: Apr. 10, 2026].",
    "MrRio, \"jsPDF documentation,\" 2025. [Online]. Available: https://github.com/parallax/jsPDF. [Accessed: Apr. 10, 2026].",
    "U. Khan, \"qrcode.react,\" 2025. [Online]. Available: https://github.com/zpao/qrcode.react. [Accessed: Apr. 10, 2026].",
    "V. Agafonkin, \"Leaflet documentation,\" 2025. [Online]. Available: https://leafletjs.com. [Accessed: Apr. 10, 2026].",
    "OpenStreetMap Foundation, \"OpenStreetMap,\" 2025. [Online]. Available: https://www.openstreetmap.org. [Accessed: Apr. 10, 2026].",
    "PostGIS Development Team, \"PostGIS documentation,\" 2025. [Online]. Available: https://postgis.net/documentation. [Accessed: Apr. 10, 2026].",
    "TypeStack, \"class-validator documentation,\" 2025. [Online]. Available: https://github.com/typestack/class-validator. [Accessed: Apr. 10, 2026].",
    "I. Seidman, Interviewing as Qualitative Research: A Guide for Researchers in Education and the Social Sciences, 4th ed. New York, NY, USA: Teachers College Press, 2013.",
]


def configure_document(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)

    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.0)


def add_title(doc: Document, chapter: str, subtitle: str) -> None:
    p1 = doc.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p1.add_run(chapter)
    r1.bold = True
    r1.font.size = Pt(16)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(subtitle)
    r2.bold = True
    r2.font.size = Pt(14)

    doc.add_paragraph("")


def add_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(13)


def add_body(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    p.paragraph_format.first_line_indent = Inches(0.3)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(6)


def add_list_item(doc: Document, text: str, numbered: bool = False) -> None:
    p = doc.add_paragraph(text)
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(3)
    if numbered:
        p.style = doc.styles["List Number"]
    else:
        p.style = doc.styles["List Bullet"]


def add_chapter_reference_note(doc: Document) -> None:
    p = doc.add_paragraph(
        "Note: All references cited in this chapter are consolidated in CHAPTER FIVE REFERENCES."
    )
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(0)
    r = p.runs[0]
    r.italic = True


def add_mermaid_code_block(doc: Document, lines: list[str]) -> None:
    for line in lines:
        p = doc.add_paragraph(line)
        p.paragraph_format.left_indent = Inches(0.4)
        p.paragraph_format.space_after = Pt(0)
        r = p.runs[0]
        r.font.name = "Courier New"
        r.font.size = Pt(10)


def build_chapter_one() -> None:
    doc = Document()
    configure_document(doc)
    add_title(
        doc,
        "CHAPTER ONE",
        "INTRODUCTION AND BACKGROUND TO THE PROJECT",
    )

    add_heading(doc, "1.0 General Introduction")
    add_body(
        doc,
        "Childhood immunization is one of the most effective public health interventions for reducing mortality and preventing severe illness in children [1], [2]. In many low-resource settings, however, the continuity of child vaccination services remains constrained by fragmented records, delayed follow-up, and weak cross-facility visibility. When records are unavailable at the point of care, children risk missed doses, delayed schedules, or unnecessary repeat vaccination.",
    )
    add_body(
        doc,
        "This project delivers a practical digital response through the Child Vaccination Command Center (CVCC), a role-based platform for tracking and managing child vaccination records across multiple healthcare branches. The system supports continuity of care, faster identity verification, timely reminders, and verifiable digital certificates while preserving data privacy and accountability [2], [4].",
    )

    add_heading(doc, "1.1 Background to the Study")
    add_body(
        doc,
        "In Ghana and related healthcare contexts, the Expanded Programme on Immunization relies heavily on paper child welfare records and branch-specific registers [4]. Although these mechanisms are familiar to health workers, they are vulnerable to physical loss, damage, incomplete updates, and delayed reporting. They also make it difficult for a nurse in one facility to view a child history first recorded in another facility.",
    )
    add_body(
        doc,
        "Global digital health guidance encourages modernized immunization systems that strengthen child-level tracking, support follow-up, and improve decision-making at facility and district levels [1], [2]. Therefore, this project explores a multi-branch digital platform that aligns frontline service workflows with managerial and public health oversight.",
    )

    add_heading(doc, "1.2 Problem Statement")
    add_body(
        doc,
        "The project problem is not the absence of immunization software globally. Rather, the practical gap is that many real deployments are either focused on aggregate reporting or are not well aligned with daily frontline workflow. In multi-branch care, this causes fragmented child histories, delayed follow-up, weak appointment visibility, and inconsistent caregiver communication [4], [8].",
    )
    add_body(
        doc,
        "The consequences include missed or overdue doses, duplicate records, manual verification delays, and reduced confidence in proof of vaccination. These weaknesses are amplified when families relocate or seek care from different facilities [1], [2], [4].",
    )

    add_heading(doc, "1.3 Aim of the Project")
    add_body(
        doc,
        "To design and implement a secure digital platform that tracks and manages child vaccination records across multiple healthcare branches while improving continuity of care for caregivers and operational efficiency for health workers.",
    )

    add_heading(doc, "1.4 Specific Objectives")
    add_list_item(
        doc,
        "Digitize child vaccination histories and make them accessible across participating branches through one centralized data layer.",
        numbered=True,
    )
    add_list_item(
        doc,
        "Implement a facility nurse workflow for search, Quick Response (QR) lookup, vaccination recording, and follow-up prioritization.",
        numbered=True,
    )
    add_list_item(
        doc,
        "Implement a Community Health Worker (CHW) outreach workflow with offline registration, offline vaccination capture, and later synchronization.",
        numbered=True,
    )
    add_list_item(
        doc,
        "Provide a caregiver portal for progress tracking, missed-dose monitoring, appointments, and certificate access.",
        numbered=True,
    )
    add_list_item(
        doc,
        "Provide verification readiness through QR-enabled digital certificates and a Public Health Authority (PHA) verification pathway.",
        numbered=True,
    )

    add_heading(doc, "1.5 Research Questions")
    add_list_item(
        doc,
        "What operational barriers do multi-branch facilities face when child vaccination records are fragmented across paper and isolated systems?",
        numbered=True,
    )
    add_list_item(
        doc,
        "How can a role-based digital platform improve child-level tracking, continuity of care, and follow-up outcomes across branches?",
        numbered=True,
    )
    add_list_item(
        doc,
        "How can QR-supported verification reduce lookup time and strengthen confidence in vaccination evidence?",
        numbered=True,
    )
    add_list_item(
        doc,
        "Which design factors most strongly influence adoption by nurses and caregivers in this context?",
        numbered=True,
    )

    add_heading(doc, "1.6 Data Collection Approach and Questionnaire Design")
    add_body(
        doc,
        "The primary data collection approach used a semi-structured questionnaire format implemented through guided interviews with nurses and selected frontline staff. The questionnaire was intentionally semi-structured so that essential questions remained consistent while still allowing respondents to explain practical workflow realities in detail [30].",
    )
    add_body(
        doc,
        "Data collection was completed in two complementary ways. First, team members called nurses by phone and asked the interview questions directly to capture real operational challenges from active practice settings. Second, the team conducted structured online research using policy documents, technical documentation, and related studies to validate observed workflow issues and identify implementation patterns [1], [2], [8].",
    )

    add_heading(doc, "1.7 Scope of the Study")
    add_body(
        doc,
        "The scope covers child vaccination record management, cross-branch retrieval, appointment support, caregiver visibility, and certificate verification workflows. The prototype does not cover adult vaccination, full national rollout governance, or vaccine supply-chain optimization in this phase.",
    )

    add_heading(doc, "1.8 Overview of the Implemented System")
    add_body(
        doc,
        "The implemented CVCC prototype supports five roles: Parent or Guardian, Facility Nurse, Community Health Worker (CHW), Data Officer, and Branch Manager. Role separation follows a Role-Based Access Control (RBAC) model and aligns each role with a dedicated interface and operational permissions [13], [20].",
    )
    add_body(
        doc,
        "At architecture level, the system follows a web client to server Application Programming Interface (API) to database model using Next.js, NestJS, and Supabase PostgreSQL [11], [13], [14], [15]. This arrangement supports centralized records with controlled access across branches.",
    )

    add_heading(doc, "1.9 Significance of the Study")
    add_body(
        doc,
        "For health workers, the system reduces lookup and verification time. For caregivers, it improves schedule visibility and proof access. For management teams, it improves branch-level monitoring and supports better intervention planning from more consistent child-level data [2], [4], [8].",
    )

    add_heading(doc, "1.10 Organization of the Study")
    add_list_item(
        doc,
        "Chapter One presents the introduction, project context, problem statement, aims, and scope.",
        numbered=True,
    )
    add_list_item(
        doc,
        "Chapter Two reviews literature, related systems, and adoption theories that support the project design.",
        numbered=True,
    )
    add_list_item(
        doc,
        "Chapter Three presents the methodology, Sprint-based development process, architecture, and system diagrams.",
        numbered=True,
    )

    add_chapter_reference_note(doc)
    doc.save(CHAPTER1_PATH)


def build_chapter_two() -> None:
    doc = Document()
    configure_document(doc)
    add_title(
        doc,
        "CHAPTER TWO",
        "LITERATURE REVIEW AND RELATED WORK",
    )

    add_heading(doc, "2.0 Introduction")
    add_body(
        doc,
        "This chapter reviews established and emerging literature on child vaccination tracking systems, multi-branch record continuity, verification mechanisms, and health technology adoption models. The review positions the Child Vaccination Command Center (CVCC) within current evidence and identifies practical gaps that the prototype addresses. It also reflects the study method used for this project, which relied mainly on online research, policy review, technical documentation, and phone-based conversations with nurses rather than formal face-to-face interviews [5], [6], [30], [31], [32].",
    )

    add_body(
        doc,
        "The purpose of the review is not only to describe what has already been written about immunization systems, but also to explain why a multi-branch digital approach is needed in the Ghanaian context. Across many settings, paper tools remain useful but limited, while digital tools improve access, continuity, and reporting when they are designed around the workflow of health workers and caregivers [1], [2], [5], [32], [35].",
    )

    add_heading(doc, "2.1 Immunization Coverage, Child Health, and the Need for Reliable Records")
    add_body(
        doc,
        "WHO reports that routine immunization remains one of the most effective public health interventions, yet coverage gaps still leave millions of children unprotected each year [1]. These gaps are not only a vaccine supply problem; they are also a record-management and follow-up problem because children can be missed when service data are incomplete, fragmented, or inaccessible across facilities [1], [2].",
    )
    add_body(
        doc,
        "The Immunization Agenda 2030 stresses that national programmes must reach every child, including those who are zero-dose or under-immunized, and that progress depends on stronger data systems, local accountability, and continuous monitoring [2]. In practical terms, this means a vaccination system must do more than store names and dates. It must help nurses, community health workers, and supervisors see who is due, who is overdue, and where follow-up is needed [1], [2], [35].",
    )
    add_body(
        doc,
        "For Ghana and similar settings, this is especially important because childhood vaccination is delivered through a mix of fixed facilities, outreach visits, and branch-level service points. When records are not connected, a child who attends a different branch may be treated as a new case, even though previous doses already exist elsewhere. That creates duplication, delays, and unreliable coverage statistics [1], [4], [35].",
    )

    add_heading(doc, "2.2 Paper-Based Immunization Registers and Their Limitations")
    add_body(
        doc,
        "Traditional child health record books and clinic registers remain common because they are inexpensive, familiar, and easy to deploy in low-resource environments [4]. However, their weaknesses are widely recognized. Paper records can be misplaced, damaged by water or wear, completed with unreadable handwriting, or simply unavailable when a child is seen at another facility [4], [31].",
    )
    add_body(
        doc,
        "In a multi-branch context, paper records also make consolidation difficult. Each branch keeps its own copy or register, but these copies are not updated automatically. As a result, reporting teams must manually combine information, which increases workload and introduces errors. This limitation matters because immunization programmes need timely and accurate data to support catch-up campaigns, stock estimation, and supervision [1], [4], [31], [32].",
    )
    add_body(
        doc,
        "Paper systems are therefore useful as backup tools, but they are not sufficient as the primary mechanism for modern multi-branch immunization management. A digital platform must preserve the convenience of frontline use while solving the access and continuity problems created by paper-only workflows [4], [31], [32].",
    )

    add_heading(doc, "2.3 Evolution of Digital Vaccination Record Keeping")
    add_body(
        doc,
        "Digital immunization systems evolved to address these limitations by introducing child-level electronic records, centralized search, automated reminders, and synchronized reporting. The main goal is to make the history of the child portable across facilities rather than trapped in a single paper folder or one clinic register [2], [32], [35].",
    )
    add_body(
        doc,
        "The global strategy on digital health emphasizes that digital initiatives must combine financial, organizational, human, and technological resources in order to succeed [31]. In other words, the presence of software alone does not guarantee better vaccination outcomes. The system must fit the health worker's routine, be supported by management, and work reliably even where infrastructure is weak [31], [32].",
    )

    add_heading(doc, "2.4 Categories of Digital Immunization Systems")
    add_heading(doc, "2.4.1 Electronic Immunization Registry (EIR) Systems")
    add_body(
        doc,
        "Electronic Immunization Registry (EIR) systems are centralized child-level records used for monitoring, follow-up, and reporting [2], [8], [35]. These systems are designed to store immunization histories in a structured form so that health workers can quickly search a child's record, determine due doses, and generate coverage reports for supervision and planning.",
    )
    add_heading(doc, "2.4.2 District Health Information Software 2 (DHIS2) Tracker Implementations")
    add_body(
        doc,
        "District Health Information Software 2 (DHIS2) Tracker implementations support individual event and longitudinal tracking modules with configurable programs [8], [35]. Tracker-based implementations are important because they allow a child to be followed over time rather than treated only as a one-time reporting entry. This makes them useful for immunization programmes that need continuous follow-up across multiple visits and service points.",
    )
    add_heading(doc, "2.4.3 Mobile Reminder Systems")
    add_body(
        doc,
        "Mobile reminder systems use Short Message Service (SMS) and app notifications to improve attendance and dose completion [1], [2], [21], [22]. These solutions support caregiver engagement by sending due-date prompts, missed-dose alerts, and follow-up messages that reduce forgetfulness and improve vaccination continuity.",
    )
    add_heading(doc, "2.4.4 Quick Response (QR) and Digital Certificate Systems")
    add_body(
        doc,
        "Quick Response (QR) and digital certificate systems provide fast identity retrieval and verifiable proof workflows [23], [24], [25]. They are useful when health workers need to confirm a child's identity quickly or when caregivers need a portable proof of vaccination that can be validated without exposing unnecessary private details.",
    )
    add_heading(doc, "2.4.5 Offline-First Mobile Capture Tools")
    add_body(
        doc,
        "Offline-first mobile capture tools allow local data entry with later synchronization to a central server when connectivity returns [16], [17], [18], [35]. These tools are essential in low-connectivity environments because they keep service delivery moving even when the internet is unstable or unavailable.",
    )

    add_body(
        doc,
        "These categories show that digital immunization systems are no longer limited to one platform type. Some solutions are designed mainly for national reporting, some for facility-level workflow, and others for caregiver communication. A strong child-vaccination platform should combine these strengths in a single architecture that supports both frontline service delivery and long-term supervision [8], [31], [35].",
    )

    add_heading(doc, "2.5 Technology Adoption Foundations")
    add_body(
        doc,
        "The Technology Acceptance Model (TAM) explains adoption through perceived usefulness and perceived ease of use [5]. In immunization settings, this implies that systems must save time and reduce complexity for nurses and support staff.",
    )
    add_body(
        doc,
        "The Unified Theory of Acceptance and Use of Technology (UTAUT) extends this perspective with performance expectancy, effort expectancy, social influence, and facilitating conditions [6]. These dimensions are highly relevant where hardware access, connectivity, and training support vary by branch.",
    )
    add_body(
        doc,
        "The Reach, Effectiveness, Adoption, Implementation, and Maintenance (RE-AIM) framework supports evaluation of scalability and long-term sustainability beyond pilot success [7].",
    )
    add_body(
        doc,
        "These theories are helpful because child vaccination systems are used by different people for different reasons. Nurses may care about speed and reliability, supervisors may care about reporting, caregivers may care about reminders and visibility, and managers may care about branch performance. A system that ignores these differences may be technically sound but practically underused [5], [6], [7], [31].",
    )

    add_heading(doc, "2.6 Electronic Immunization Registries and DHIS2 Tracker")
    add_body(
        doc,
        "Electronic Immunization Registries (EIRs) are a core response to the limitations of paper-based vaccination records. They make it possible to store child-level data centrally, search quickly by identity, follow vaccination history over time, and generate reports for districts and ministries [2], [8], [35].",
    )
    add_body(
        doc,
        "DHIS2 Tracker is especially relevant because it bridges individual-level tracking and broader health information management. The Tracker and Event model supports longitudinal records, mobile capture, duplicate search, automatic aggregation, appointment follow-up, and offline data entry with later synchronization [8], [35]. These features directly match the workflow needs of a multi-branch child vaccination system.",
    )
    add_body(
        doc,
        "In practice, DHIS2 demonstrates an important design lesson: the best immunization systems are not just databases. They are workflow tools. They help health workers register a child once, follow the child across visits, reduce duplicate entry, and support both local service delivery and higher-level reporting without forcing staff to re-enter the same information in many places [8], [35].",
    )

    add_heading(doc, "2.7 Mobile Reminder Systems, Caregiver Engagement, and Follow-Up")
    add_body(
        doc,
        "A recurring theme in immunization literature is that missed appointments are often not caused by refusal alone. Some doses are missed because caregivers forget dates, relocate, misunderstand schedules, or never receive a follow-up message in time [1], [2], [21], [22].",
    )
    add_body(
        doc,
        "Digital reminder systems address this gap through SMS alerts, mobile app notifications, and structured follow-up messaging. These tools are especially useful in settings where mobile phone access is widespread but internet access may still be uneven [21], [22]. When reminders are tied to actual due dates, vaccination attendance improves and the workload of nurses is reduced because fewer children are missed silently [1], [2].",
    )
    add_body(
        doc,
        "For this project, reminder support is important because the platform is not meant to be a passive register. It is designed to help the parent, the nurse, and the CHW all see the same child journey in a structured way. A caregiver who can see the next vaccine date, a missed dose, or a downloadable certificate is more likely to engage with the immunization process consistently [21], [22], [24], [25].",
    )

    add_heading(doc, "2.8 Security, Privacy, Authentication, and Trust")
    add_body(
        doc,
        "Child health records are sensitive and require controlled access, secure authentication, and traceable actions. Secure system design commonly combines Role-Based Access Control (RBAC), encrypted transport, and controlled token handling such as Json Web Token (JWT) session enforcement [19], [20].",
    )
    add_body(
        doc,
        "Trust in vaccination evidence is strengthened when certificate verification confirms authenticity without unnecessary disclosure of private data [19], [20], [23], [24], [25]. That means the system should show enough information for verification while protecting the caregiver's and child's privacy. In a practical deployment, this includes role-based permissions, secure credentials, and auditability of sensitive actions [19], [20], [31].",
    )
    add_body(
        doc,
        "Security is also part of usability. If staff constantly struggle with logins, timeouts, or unclear permissions, they may bypass the system or revert to paper. Therefore, the literature supports security designs that are strong but not unnecessarily complicated for frontline users [19], [20], [31], [32].",
    )

    add_heading(doc, "2.9 Offline-First Design and Low-Connectivity Settings")
    add_body(
        doc,
        "Connectivity limitations are common in outreach and rural service settings, so a child vaccination system must not depend entirely on a live network connection. Offline-first mobile design allows data to be captured locally, validated on the device, and synchronized later when a connection becomes available [16], [17], [18], [35].",
    )
    add_body(
        doc,
        "This approach is especially relevant for Community Health Workers, who may visit homes, schools, or temporary outreach sites where network availability is poor. If the system can store the record safely at the point of service, the worker can still perform the core task without interruption, while the platform later aligns the record with the central branch database [16], [17], [18].",
    )
    add_body(
        doc,
        "DHIS2 Tracker documents a similar principle: mobile capture can work in both online and offline modes and automatically sync when internet access returns [35]. The same logic supports the CVCC CHW workflow, where an offline register can later be reconciled with the branch-level and central record store.",
    )

    add_heading(doc, "2.10 Ghanaian Context and Research Method Alignment")
    add_body(
        doc,
        "This project was grounded in online research and phone-based contact with nurses rather than formal face-to-face interviews. That approach was appropriate because the project needed timely practical insights from busy healthcare staff while still validating the design against current guidance and technical evidence [30], [31], [32].",
    )
    add_body(
        doc,
        "The Ghanaian context makes this especially important. Health facilities may vary in infrastructure, staffing, and workflow maturity, which means a system that works in one branch may fail in another if it is too dependent on fixed connectivity or excessive administrative steps. The literature therefore supports a pragmatic, workflow-aware, and branch-synchronized solution rather than a one-size-fits-all tool [4], [31], [35].",
    )
    add_body(
        doc,
        "The current system analysis in the project confirms that paper registers and isolated records still dominate many service points. This creates a real research gap that the CVCC prototype is designed to address: a single child record that can move across branches, remain secure, and still support supervisory reporting and caregiver communication [4], [8], [31], [35].",
    )

    add_heading(doc, "2.11 Research Gap and Positioning of CVCC")
    add_body(
        doc,
        "The literature confirms that digital immunization tools exist, but common implementation gaps remain in integrated frontline usability, caregiver self-service, and cross-branch continuity. Many systems are either reporting-centric or incomplete for routine nurse and Community Health Worker (CHW) operations [1], [2], [8], [35].",
    )
    add_body(
        doc,
        "The CVCC prototype addresses this gap by combining facility workflow support, caregiver tracking, QR-based verification, and offline outreach synchronization in one coordinated platform. It aligns operational and supervisory needs while remaining realistic for incremental deployment.",
    )
    add_body(
        doc,
        "In this sense, the project is not trying to replace existing public health tools. Instead, it aims to strengthen the missing layer between paper-based service delivery and aggregated reporting systems by providing child-level continuity across branches, roles, and service points [2], [8], [31], [35].",
    )

    add_heading(doc, "2.12 Comparative Review of Existing Systems and Practical Gaps")
    add_body(
        doc,
        "A comparison of the reviewed literature shows that many existing digital health tools solve only part of the child immunization problem. Some systems focus on registration and reporting, while others focus on reminders, certificate issuance, or verification. A smaller number support individual tracking, but even those may not fully address branch-to-branch continuity, offline outreach usage, or caregiver-facing self-service in one integrated workflow [2], [7], [8], [9], [13], [35].",
    )
    add_body(
        doc,
        "This partial coverage is important because immunization is not a single event. It is a sequence of interactions between the child, the caregiver, the nurse, the community health worker, and the supervisor. If a system only records the current visit but does not preserve the history across branches, then the child still risks duplication, missed follow-up, and poor continuity of care. If a system only supports reminders but not reliable search and verification, then caregivers may be informed but still unable to prove status or recover old records [1], [2], [8], [21], [22].",
    )
    add_body(
        doc,
        "The literature also suggests that the usability of a system matters as much as its technical completeness. A highly capable registry that is slow to load, hard to navigate, or difficult to use under clinic pressure is unlikely to be used consistently. TAM and UTAUT both support this argument by showing that perceived usefulness, ease of use, performance expectancy, and facilitating conditions influence whether health workers adopt new digital tools [5], [6].",
    )
    add_body(
        doc,
        "This is one reason the CVCC design intentionally combines a limited number of clear role-specific dashboards with structured workflows such as search, registration, reminders, and verification. The goal is not feature overload; the goal is practical completeness. The literature consistently favors systems that reduce manual workload, fit the everyday environment of health workers, and preserve child-level continuity at the point of care [4], [8], [31], [35].",
    )

    add_heading(doc, "2.13 Synthesis of the Literature and Design Implications for CVCC")
    add_body(
        doc,
        "When the reviewed sources are synthesized together, five design implications stand out. First, the system must maintain a single child record that can be searched, updated, and verified across branches without duplication. Second, it must support caregivers through reminders and clear appointment visibility. Third, it must tolerate poor connectivity by allowing offline capture and later synchronization. Fourth, it must protect sensitive child data through secure authentication and role-based access. Fifth, it must support branch, district, and public health reporting without forcing duplicate manual entry [1], [2], [8], [19], [20], [21], [22], [31], [35].",
    )
    add_body(
        doc,
        "These implications directly shaped the CVCC prototype. The parent module focuses on visibility and proof access. The facility module focuses on search, recording, and verification. The CHW module focuses on offline-first outreach workflows. The branch, headquarters, data officer, and PHA modules focus on monitoring, supervision, deduplication, and reporting. This mapping is consistent with the literature's emphasis on using digital systems to strengthen existing workflows rather than replace them abruptly [5], [6], [7], [8], [31], [35].",
    )
    add_body(
        doc,
        "A further implication is that health systems in Ghana and similar environments need adaptable tools rather than rigid one-size-fits-all packages. A branch may have stable internet today but intermittent connectivity tomorrow; a nurse may need quick lookup during peak hours; a CHW may need to capture a record in a remote community; and a parent may need a reminder on a basic phone. The literature therefore supports an architecture that is flexible in delivery but strict in data integrity [16], [17], [18], [19], [20], [31], [35].",
    )

    add_heading(doc, "2.14 Chapter Summary")
    add_body(
        doc,
        "Chapter Two establishes the theoretical and practical basis for the project. It confirms the importance of child-level continuity, secure verification, adoption-aware design, digital reminder support, and offline-first capture. These findings directly inform the methodology and implementation decisions documented in Chapter Three.",
    )
    add_body(
        doc,
        "Overall, the review shows that a useful child vaccination platform must be built around both evidence and workflow. It must answer the needs of nurses, Community Health Workers, parents, branch managers, data officers, headquarters administrators, and public health reviewers while still remaining practical in low-connectivity environments [5], [6], [7], [31], [35].",
    )

    add_chapter_reference_note(doc)
    doc.save(CHAPTER2_PATH)


def build_chapter_three() -> None:
    doc = Document()
    configure_document(doc)
    add_title(
        doc,
        "CHAPTER THREE",
        "METHODOLOGY AND SYSTEM DESIGN",
    )

    add_heading(doc, "3.0 Introduction")
    add_body(
        doc,
        "This chapter explains how the Child Vaccination Command Center (CVCC) was designed and implemented from requirements discovery through architecture, testing, and documentation. The project is organized around five user roles: Parent or Guardian, Facility Nurse, Community Health Worker (CHW), Data Officer, and Branch Manager. The chapter emphasizes methodological decisions, Sprint execution, and practical constraints from frontline healthcare operations.",
    )

    add_heading(doc, "3.1 Research Methodology and Data Collection")
    add_body(
        doc,
        "The study used a semi-structured questionnaire design implemented through guided interviews with nurses and selected field staff [30]. The semi-structured approach ensured question consistency while allowing respondents to explain branch-specific workflow realities, including follow-up behavior, record transfer difficulties, and verification bottlenecks.",
    )
    add_body(
        doc,
        "Primary collection was performed through direct phone calls to nurses. Secondary evidence was gathered through online research from public health policy documents, platform documentation, and peer-reviewed methods sources [1], [2], [8]. This dual method reduced bias and improved requirement completeness.",
    )
    add_body(
        doc,
        "Key constraints identified from data collection included intermittent connectivity in outreach locations, inconsistent cross-branch child transfer handling, and slow certificate verification workflows. These constraints directly informed the technical architecture and module priorities.",
    )

    add_heading(doc, "3.2 Software Development Methodology With Sprint Model")
    add_body(
        doc,
        "Development followed Agile principles with a Scrum-aligned Sprint process [9], [10]. Work was organized into short iterative cycles, and each Sprint delivered a testable increment of value for one or more user roles.",
    )
    add_list_item(
        doc,
        "Sprint Planning: define user stories, acceptance criteria, and module dependencies.",
    )
    add_list_item(
        doc,
        "Sprint Execution: implement frontend and backend slices with daily synchronization by the team.",
    )
    add_list_item(
        doc,
        "Sprint Review: demonstrate completed functionality against requirements and gather feedback.",
    )
    add_list_item(
        doc,
        "Sprint Retrospective: identify blockers, quality gaps, and process improvements for the next cycle.",
    )
    add_body(
        doc,
        "This Sprint cadence enabled controlled scope evolution and quick correction of defects while preserving shared ownership across team members.",
    )

    add_heading(doc, "3.3 Sprint Workflow Diagram in Mermaid Code")
    add_body(
        doc,
        "The project Sprint cycle is represented below in Mermaid code form so it can be reused directly in documentation tooling or diagram rendering engines.",
    )
    add_mermaid_code_block(
        doc,
        [
            "flowchart LR",
            "    A[Product Backlog - all pending tasks] --> B[Sprint Planning - choose tasks for this cycle]",
            "    B --> C[Sprint Backlog - selected tasks for this Sprint]",
            "    C --> D[Sprint Execution - build and test the selected tasks]",
            "    D --> E[Daily Standup - quick team progress check]",
            "    E --> D",
            "    D --> F[Sprint Review - demonstrate completed work]",
            "    F --> G[Sprint Retrospective - discuss improvements for next Sprint]",
            "    G --> H[Increment Release - deliver working features]",
            "    H --> A",
        ],
    )

    add_heading(doc, "3.4 System Architecture and Technology Stack")
    add_body(
        doc,
        "The platform uses a three-tier model: web client, secure backend Application Programming Interface (API), and centralized relational database [11], [13], [14], [15]. The frontend is implemented with Next.js and React, while backend services are implemented in NestJS with typed request validation [11], [12], [13], [29].",
    )
    add_body(
        doc,
        "Security controls include Role-Based Access Control (RBAC), Json Web Token (JWT) based session enforcement, and controlled data operations over server-side services [13], [20]. In the database layer, access controls and operational logging support accountability across user roles.",
    )
    add_body(
        doc,
        "For outreach resilience, the Community Health Worker (CHW) module uses an offline-first approach with Indexed Database (IndexedDB) storage through Dexie.js and Progressive Web Application (PWA) service-worker caching [16], [17], [18]. Sensitive local data handling aligns with Advanced Encryption Standard (AES) principles [19].",
    )

    add_heading(doc, "3.5 Core Modules Implemented")
    add_list_item(
        doc,
        "Identity and lookup: child search and Quick Response (QR) scanning workflows [23], [25].",
    )
    add_list_item(
        doc,
        "Vaccination capture: facility and outreach record entry with synchronization and conflict handling.",
    )
    add_list_item(
        doc,
        "Certificates: Portable Document Format (PDF) generation and verification flow [24], [25].",
    )
    add_list_item(
        doc,
        "Notifications: Short Message Service (SMS) and email reminders for appointments and due doses [21], [22].",
    )
    add_list_item(
        doc,
        "Monitoring: role-specific dashboards for branch managers and data officers.",
    )

    add_heading(doc, "3.6 Testing and Quality Assurance")
    add_body(
        doc,
        "Testing combined functional walkthroughs, endpoint validation, and offline synchronization verification. Integration checks confirmed request validation behavior, role-guard enforcement, and expected data consistency after sync. Offline tests simulated network loss and restoration to verify queue integrity and idempotent replay behavior [16], [17], [29].",
    )
    add_body(
        doc,
        "Defects discovered during implementation were documented and resolved iteratively within Sprint reviews and retrospectives, improving reliability before subsequent module releases.",
    )

    add_heading(doc, "3.7 System Diagrams")
    add_body(
        doc,
        "This section presents all system diagrams supplied for the project. Each diagram is inserted as evidence of system modeling, process flow, and role interaction design.",
    )

    diagram_files = sorted(DIAGRAMS_DIR.glob("*.png"))
    if not diagram_files:
        add_body(doc, "No diagram image was found in the configured diagrams folder at generation time.")
    else:
        for idx, img in enumerate(diagram_files, start=1):
            caption_text = img.stem.replace("-", " ").replace("_", " ").strip().title()
            add_heading(doc, f"3.7.{idx} {caption_text}")
            pic_paragraph = doc.add_paragraph()
            pic_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = pic_paragraph.add_run()
            run.add_picture(str(img), width=Inches(6.0))

    add_heading(doc, "3.8 Chapter Summary")
    add_body(
        doc,
        "Chapter Three documented the project methodology, Sprint-driven development approach, architecture decisions, quality strategy, and system modeling artifacts. The chapter demonstrates that implementation decisions were directly linked to field evidence from semi-structured interviews and online research, and that the final system design aligns with the five-user scope of the project.",
    )

    add_chapter_reference_note(doc)
    doc.save(CHAPTER3_PATH)


def build_references_doc() -> None:
    doc = Document()
    configure_document(doc)

    add_title(doc, "CHAPTER FIVE", "REFERENCES")
    add_body(
        doc,
        "This chapter consolidates all references cited across Chapter One, Chapter Two, and Chapter Three using Institute of Electrical and Electronics Engineers (IEEE) citation formatting.",
    )

    for index, item in enumerate(REFERENCES, start=1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.4)
        p.paragraph_format.first_line_indent = Inches(-0.4)
        p.paragraph_format.space_after = Pt(5)
        p.paragraph_format.line_spacing = 1.3

        num_run = p.add_run(f"[{index}] ")
        num_run.bold = True
        text_run = p.add_run(item)
        text_run.bold = False

    doc.save(REFERENCES_PATH)


def main() -> None:
    build_chapter_one()
    build_chapter_two()
    build_chapter_three()
    build_references_doc()

    print(f"Generated: {CHAPTER1_PATH}")
    print(f"Generated: {CHAPTER2_PATH}")
    print(f"Generated: {CHAPTER3_PATH}")
    print(f"Generated: {REFERENCES_PATH}")


if __name__ == "__main__":
    main()
