"""
Generates CHAPTER-3.docx with proper IEEE academic formatting
and numbered IEEE-style references embedded inline and listed at end.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ─── helpers ─────────────────────────────────────────────────────────────────

def set_paragraph_spacing(para, before=0, after=6, line_rule=WD_LINE_SPACING.ONE_POINT_FIVE):
    pf = para.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after  = Pt(after)
    pf.line_spacing_rule = line_rule

def set_font(run, name="Times New Roman", size=12, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def body_para(doc, text_parts, indent=False):
    """
    text_parts: list of (text, bold, italic, mono)
    A normal body paragraph composed of formatted runs.
    """
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=0, after=6)
    pf = p.paragraph_format
    pf.first_line_indent = Inches(0.3) if indent else Inches(0)
    if indent:
        pf.left_indent = Inches(0)
    for (text, bold, italic, mono) in text_parts:
        r = p.add_run(text)
        if mono:
            r.font.name = "Courier New"
            r.font.size = Pt(10)
        else:
            r.font.name = "Times New Roman"
            r.font.size = Pt(12)
        r.font.bold = bold
        r.font.italic = italic
    return p

def plain(t):
    return (t, False, False, False)

def bold(t):
    return (t, True, False, False)

def italic(t):
    return (t, False, True, False)

def code(t):
    return (t, False, True, True)

def cite(numbers):
    """Return a plain run for an IEEE citation, e.g. [1] or [3], [4]"""
    tag = "[" + ", ".join(str(n) for n in numbers) + "]"
    return (tag, False, False, False)

def add_heading(doc, text, level):
    """
    level 0 = chapter title (centre, 16pt bold)
    level 1 = 3.x  (14pt bold)
    level 2 = 3.x.x (13pt bold)
    level 3 = 3.x.x.x (12pt bold italic) — not used here but available
    """
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=12, after=4)
    r = p.add_run(text)
    if level == 0:
        r.font.name = "Times New Roman"; r.font.size = Pt(16); r.font.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
    elif level == 1:
        r.font.name = "Times New Roman"; r.font.size = Pt(14); r.font.bold = True
        p.paragraph_format.space_before = Pt(18)
    elif level == 2:
        r.font.name = "Times New Roman"; r.font.size = Pt(13); r.font.bold = True
        p.paragraph_format.space_before = Pt(12)
    return p

def add_spacer(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)

# ─── document setup ──────────────────────────────────────────────────────────

doc = Document()

# Page margins (2.5 cm all round is common for Ghanaian universities)
for section in doc.sections:
    section.top_margin    = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin   = Inches(1.25)
    section.right_margin  = Inches(1.0)

# ─── CHAPTER TITLE ───────────────────────────────────────────────────────────
add_heading(doc, "CHAPTER THREE", 0)
add_heading(doc, "METHODOLOGY AND SYSTEM DESIGN", 0)
add_spacer(doc)

# ─── 3.0 INTRODUCTION ────────────────────────────────────────────────────────
add_heading(doc, "3.0  Introduction", 1)

body_para(doc, [
    plain("This chapter presents a detailed account of the methodology, system architecture, "
          "and technology choices that guided the development of the Child Vaccination "
          "Coordination Centre (CVCC) system. Building a reliable vaccination management "
          "platform for Ghana's primary healthcare environment demanded far more than simply "
          "selecting the most popular frameworks; it required a careful reconciliation of "
          "real-world clinical constraints—such as low Internet connectivity in rural "
          "communities, cross-facility child transfers, and strict patient data privacy—with "
          "the capabilities afforded by modern web technologies "),
    cite([3]),
    plain("."),
], indent=True)

body_para(doc, [
    plain("The chapter begins by describing how data was collected from nurses and how the "
          "insights uncovered during that process shaped the fundamental design decisions of "
          "the system. It then explains the Agile software development methodology adopted by "
          "the three-member team and how work was coordinated across seven distinct user roles "),
    cite([1]),
    plain(". The chapter proceeds to describe the system's three-tier architecture, with "
          "particular attention to the offline-first Community Health Worker (CHW) module, "
          "which is perhaps the most technically demanding aspect of the entire project. The "
          "chapter closes with sections on the specific technology stack chosen for each layer "
          "of the system, the core modules and how they were implemented, and the testing and "
          "quality assurance strategies employed throughout development."),
], indent=True)

# ─── 3.1 ─────────────────────────────────────────────────────────────────────
add_heading(doc, "3.1  Research Methodology and Data Collection", 1)

add_heading(doc, "3.1.1  Data Collection Methods", 2)

body_para(doc, [
    plain("The CVCC system was not designed in isolation. From the very beginning, the team "
          "recognised that the most significant risk of failure was building a technically "
          "impressive system that failed to reflect how healthcare workers actually operate on "
          "the ground. To avoid this, primary data was collected through semi-structured "
          "interviews conducted with nurses and community health workers at selected healthcare "
          "facilities "),
    cite([2]),
    plain(". This format was deliberately chosen over rigid structured questionnaires because "
          "it allowed conversations to flow naturally—nurses were encouraged to describe their "
          "day-to-day workflows, the frustrations they experience with the current paper-based "
          "system, and the specific scenarios that most often lead to errors or data loss."),
], indent=True)

body_para(doc, [
    plain("Through these interviews, several recurring pain points were identified. First, "
          "nurses frequently had no way of confirming whether a child brought to their facility "
          "had already received certain vaccines at another facility. This problem was especially "
          "common in contexts where families had recently relocated or were receiving care across "
          "different healthcare centres. Second, the process of issuing and verifying vaccination "
          "certificates was entirely manual; there was no mechanism by which a certificate's "
          "authenticity could be independently confirmed. Third, parents and guardians were not "
          "routinely notified of upcoming vaccination appointments, meaning that missed doses "
          "were often discovered only after they had already passed "),
    cite([3]),
    plain("."),
], indent=True)

body_para(doc, [
    plain("To complement the interview findings, a thorough literature review was also "
          "conducted. This review covered published works on immunisation information systems "
          "in sub-Saharan Africa, WHO and UNICEF guidelines on digital health implementations "),
    cite([3]),
    plain(", and existing vaccination registry platforms deployed in other developing countries. "
          "The synthesis of interview results and literature findings gave the team a "
          "well-rounded understanding of both the local operational context and the broader "
          "global landscape of child immunisation data management."),
], indent=True)

add_heading(doc, "3.1.2  Analysis of Field Constraints", 2)

body_para(doc, [
    plain("The field research surfaced three constraints that had a direct and lasting impact "
          "on the system's architecture. The first was "),
    bold("network unavailability"),
    plain(". A significant portion of the children served by community health workers live in "
          "areas where mobile data connectivity is either intermittent or entirely absent. Any "
          "system that relied on a constant Internet connection for the CHW would be practically "
          "unusable in these communities. This insight made an offline-first design "),
    cite([18]),
    plain(" not merely desirable but an absolute requirement."),
], indent=True)

body_para(doc, [
    plain("The second constraint was "),
    bold("catchment area rules"),
    plain(". Ghana's primary healthcare delivery is organised around catchment areas—geographic "
          "zones assigned to specific healthcare facilities. A child is typically expected to "
          "receive vaccinations at the facility that serves their residential area. However, "
          "fieldwork revealed that this was not always followed in practice; families regularly "
          "moved between districts, visited relatives in other areas, or simply had more "
          "convenient access to a facility outside their assigned catchment. The system therefore "
          "had to support the concept of a child transferring between facilities in a structured, "
          "auditable way."),
], indent=True)

body_para(doc, [
    plain("The third constraint was "),
    bold("patient transfer scenarios"),
    plain(". Connected to catchment rules, it was clear that the system needed explicit workflows "
          "to handle transfer-in (accepting a child from another healthcare facility's register) "
          "and transfer-out (releasing a child to another facility's register) operations. This "
          "had to be more than a simple data update; it needed to preserve the full vaccination "
          "history of the child regardless of which facility was managing their care at any "
          "given point in time."),
], indent=True)

# ─── 3.2 ─────────────────────────────────────────────────────────────────────
add_heading(doc, "3.2  Software Development Life Cycle (SDLC)", 1)

add_heading(doc, "3.2.1  Agile Methodology Adoption", 2)

body_para(doc, [
    plain("The development of the CVCC system followed an Agile software development "
          "methodology "),
    cite([1]),
    plain(", specifically drawing from the principles of iterative and incremental delivery. "
          "Rather than attempting to design the complete system on paper before writing a single "
          "line of code, the team adopted a slice-by-slice approach in which each user role was "
          "treated as a vertical slice of functionality that would be built, tested, and "
          "demonstrated before moving on to the next."),
], indent=True)

body_para(doc, [
    plain("This decision was especially well-suited to the nature of the project. The system "
          "serves seven distinct user roles—HQ Administrator, Branch Manager, Facility Nurse, "
          "Community Health Worker (CHW), Data Officer, Public Health Analyst (PHA), and "
          "Parent—and each role has a materially different set of responsibilities, views, and "
          "permissions. Attempting to design all of them simultaneously before any implementation "
          "was underway would have introduced significant confusion and wasted effort, particularly "
          "given that insights from building one role often revealed dependencies or shared "
          "components useful to another."),
], indent=True)

body_para(doc, [
    plain("Working iteratively also made it easier to incorporate feedback at regular intervals. "
          "After completing a functional slice for one role, the team could review it, identify "
          "weaknesses, and carry those learnings into the next iteration. For example, the "
          "reusable "),
    code("shadcn/ui"),
    plain(" component library "),
    cite([8]),
    plain(" and the shared API configuration in "),
    code("lib/api/config.ts"),
    plain(" were both outcomes of decisions made after building the first few role dashboards, "
          "when the team recognised the need for consistency across the interface."),
], indent=True)

add_heading(doc, "3.2.2  Team Coordination and Version Control", 2)

body_para(doc, [
    plain("The project was built by a team of three developers. To prevent duplication of "
          "effort and conflicting code changes, the seven user roles were divided across the "
          "team members, with each developer taking primary responsibility for both the frontend "
          "dashboard and the backend module of the roles assigned to them. This division ensured "
          "that at any given point in time, no two developers were working on the same files "
          "simultaneously."),
], indent=True)

body_para(doc, [
    plain("The codebase was maintained in a single Git repository, with the Next.js frontend "
          "at the root of the repository and the NestJS backend housed in a dedicated "),
    code("backend/"),
    plain(" subdirectory. This monorepo-style layout meant that both the frontend and the "
          "backend shared the same version history, making it straightforward to review any "
          "change to either layer. Developers worked on feature branches, submitted pull "
          "requests for review, and merged only after the changes had been examined by at "
          "least one other team member. This process enforced a minimum level of code review "
          "on every non-trivial change and made it possible to roll back any breaking change "
          "cleanly. A "),
    code("CHANGELOG.md"),
    plain(" file was maintained throughout development to document significant changes and the "
          "reasoning behind key decisions."),
], indent=True)

# ─── 3.3 ─────────────────────────────────────────────────────────────────────
add_heading(doc, "3.3  System Architecture", 1)

add_heading(doc, "3.3.1  High-Level Architecture Overview", 2)

body_para(doc, [
    plain("The CVCC system is built on a three-tier architecture consisting of a Progressive "
          "Web Application (PWA) "),
    cite([18]),
    plain(" frontend, a RESTful API backend, and a cloud-hosted relational database. At the "
          "frontend layer, a "),
    bold("Next.js 16"),
    plain(" application "),
    cite([4]),
    plain(" built with the App Router paradigm serves as the user interface for all seven "
          "roles. At the backend layer, a "),
    bold("NestJS 10"),
    plain(" application "),
    cite([6]),
    plain(" exposes a structured set of REST API endpoints that the frontend calls for all "
          "data operations. At the data layer, a "),
    bold("Supabase-hosted PostgreSQL"),
    plain(" database "),
    cite([14]),
    plain(" stores all persistent records, with Row-Level Security (RLS) policies enforcing "
          "data access rules at the database level."),
], indent=True)

body_para(doc, [
    plain("The flow of a typical interaction—for example, a facility nurse recording a vaccine "
          "administration—begins at the React component "),
    cite([5]),
    plain(" in the browser, which calls a typed API function in "),
    code("lib/api/"),
    plain(". This function sends an authenticated HTTP request (with a JWT token "),
    cite([21]),
    plain(" carried in an HttpOnly cookie) to the relevant NestJS endpoint, such as "),
    code("POST /api/facility/children/:id/vaccinations"),
    plain(". The NestJS controller validates the request body using "),
    code("class-validator"),
    plain(" DTOs "),
    cite([24]),
    plain(", applies the "),
    code("@Roles('facility-nurse')"),
    plain(" guard to confirm authorisation, and delegates to the corresponding service method. "
          "The service method then calls the "),
    code("DatabaseService"),
    plain(", which uses the Supabase JavaScript client with a service-role key to perform the "
          "database write and return the result up the chain to the frontend."),
], indent=True)

body_para(doc, [
    plain("This separation of concerns—frontend calls API, API calls database—was a deliberate "
          "architectural choice. It means the frontend never holds any privileged database "
          "credentials, all business logic lives in one authoritative location (the NestJS "
          "services), and the database access layer can be independently modified without "
          "touching the UI."),
], indent=True)

add_heading(doc, "3.3.2  Offline-First Architecture (The CHW Module)", 2)

body_para(doc, [
    plain("The CHW module is architecturally distinct from every other role in the system, "
          "and it represents the most technically complex component of the CVCC platform. "
          "Because community health workers conduct home visits and outreach sessions in areas "
          "with little or no Internet connectivity, the system must continue to function fully "
          "even when the device has no network access. This requirement demanded an offline-first "
          "design "),
    cite([18]),
    plain(" at every layer of the CHW module."),
], indent=True)

body_para(doc, [
    plain("The offline data layer is built on "),
    bold("Dexie.js"),
    plain(" "),
    cite([10]),
    plain(", a typed wrapper around the browser's IndexedDB API "),
    cite([19]),
    plain(", instantiated as the "),
    code("cvcc_chw_offline_v2"),
    plain(" database. This local database contains three logical tables: "),
    code("children"),
    plain(" (indexed by "),
    code("cvccId"),
    plain(" and "),
    code("catchmentAreaId"),
    plain("), "),
    code("vaccinationQueue"),
    plain(" (the write-ahead action log), and "),
    code("offlineMapStatus"),
    plain(" (tracking the state of cached map tiles). When a CHW registers a child, records "
          "a vaccination, or initiates a transfer while offline, the action—along with all its "
          "data—is written into the "),
    code("vaccinationQueue"),
    plain(" as a "),
    code("VaccinationQueueItem"),
    plain(". This item carries an "),
    code("idempotencyKey"),
    plain(" (generated from the child ID, vaccine ID, and timestamp) to guarantee that the "
          "same action is never processed twice on the server, even if the sync happens more "
          "than once."),
], indent=True)

body_para(doc, [
    plain("Synchronisation is handled by the "),
    code("ChwBackgroundSyncService"),
    plain(" in "),
    code("lib/chw-offline/background-sync.ts"),
    plain(". This singleton service runs on a five-minute polling interval and also listens "
          "for the browser's "),
    code("window.online"),
    plain(" event, so that when the device reconnects to a network, it immediately begins "
          "draining the pending queue "),
    cite([25]),
    plain(". It calls the "),
    code("POST /api/chw/vaccinations/sync"),
    plain(" endpoint to deliver the batched actions to the NestJS backend, which processes "
          "each item and returns a confirmation."),
], indent=True)

body_para(doc, [
    plain("Two separate service workers augment the Dexie-based solution. The primary CHW "
          "service worker ("),
    code("public/chw-service-worker.js"),
    plain(") uses a network-first strategy for all CHW pages and API routes, falling back to "
          "a cache named "),
    code("chw-offline-cache-v1"),
    plain(" when the network is unavailable. It also registers a "),
    code("sync-chw-vaccinations"),
    plain(" background sync tag "),
    cite([25]),
    plain(" so that the browser itself can attempt re-delivery when connectivity is restored. "
          "A second dedicated service worker ("),
    code("public/chw-map-sw.js"),
    plain(") uses a cache-first strategy specifically for OpenStreetMap "),
    cite([12]),
    plain(" tile images, caching them under the name "),
    code("cvcc-chw-osm-tiles-v1"),
    plain(". This means that once a CHW has explored a geographic area while online, the map "
          "tiles for that area remain available for subsequent offline visits."),
], indent=True)

body_para(doc, [
    plain("Because the CHW local database stores personally identifiable information (PII) "
          "about children and their guardians, all sensitive fields—including "),
    code("fullName"),
    plain(", "),
    code("dateOfBirth"),
    plain(", "),
    code("guardianName"),
    plain(", and "),
    code("guardianPhone"),
    plain("—are encrypted before being committed to IndexedDB. The encryption system in "),
    code("lib/chw-offline/encryption.ts"),
    plain(" uses AES-GCM-256 "),
    cite([20]),
    plain(" via the browser's native Web Crypto API, with an encryption key derived from the "
          "authenticated user's ID and the first 32 characters of their access token. This key "
          "is held only in "),
    code("sessionStorage"),
    plain(" and is never persisted to disk. Additional security measures in the CHW module "
          "include an auto-logout timer at 15 minutes of inactivity ("),
    code("lib/chw-offline/auto-logout.ts"),
    plain(") and an auto-clear policy that wipes all IndexedDB data after seven days of device "
          "inactivity ("),
    code("lib/chw-offline/auto-clear.ts"),
    plain("). A client-side audit log in "),
    code("localStorage"),
    plain(" records all read and write operations for the current session, providing a "
          "tamper-evident trail of what the CHW accessed while offline."),
], indent=True)

add_heading(doc, "3.3.3  Role-Based Access Control (RBAC) Architecture", 2)

body_para(doc, [
    plain("The CVCC system manages seven distinct user roles: "),
    code("hq-admin"),
    plain(", "),
    code("branch-manager"),
    plain(", "),
    code("facility-nurse"),
    plain(", "),
    code("chw"),
    plain(", "),
    code("data-officer"),
    plain(", "),
    code("pha"),
    plain(", and "),
    code("parent"),
    plain(". Each role has a fundamentally different scope of operations and must not be able "
          "to access data or functionality intended for another role. This isolation is enforced "
          "at three distinct layers of the architecture."),
], indent=True)

body_para(doc, [
    plain("At the "),
    bold("transport layer"),
    plain(", all authenticated requests carry a JWT token "),
    cite([21]),
    plain(" issued by the NestJS "),
    code("AuthService"),
    plain(". The token is stored in an HttpOnly cookie with a seven-day expiry, making it "
          "inaccessible to any JavaScript running on the page. The "),
    code("JwtStrategy"),
    plain(" in "),
    code("backend/src/auth/strategies/jwt.strategy.ts"),
    plain(" validates every incoming request and attaches the decoded user profile—including "
          "their role—to the request object."),
], indent=True)

body_para(doc, [
    plain("At the "),
    bold("API layer"),
    plain(", every NestJS controller that serves role-specific data is protected by both the "),
    code("JwtAuthGuard"),
    plain(" (confirming that the request is authenticated) and the "),
    code("RolesGuard"),
    plain(" (confirming that the authenticated user holds the correct role). The "),
    code("@Roles()"),
    plain(" decorator is applied at the controller or handler level: for instance, the entire "
          "CHW controller is annotated with "),
    code("@Roles('chw')"),
    plain(", the PHA controller with "),
    code("@Roles('pha')"),
    plain(", and so on. Any request arriving at a guarded endpoint with an incorrect role "
          "receives a 403 Forbidden response before any business logic is executed."),
], indent=True)

body_para(doc, [
    plain("At the "),
    bold("database layer"),
    plain(", Supabase Row-Level Security (RLS) policies "),
    cite([14]),
    plain(" provide a final defence-in-depth layer. Even if a request were to somehow bypass "
          "the API guards, the RLS policies on each table restrict which rows a given user can "
          "read or modify based on their role and their associated "),
    code("branch_id"),
    plain(". The NestJS "),
    code("DatabaseService"),
    plain(" in "),
    code("backend/src/common/database/database.service.ts"),
    plain(" uses the Supabase service-role key, which is intentionally held only on the server "
          "and never exposed to the frontend, ensuring that all database interactions are "
          "mediated through the API's authorisation logic."),
], indent=True)

# ─── 3.4 ─────────────────────────────────────────────────────────────────────
add_heading(doc, "3.4  Technology Stack Selection", 1)

add_heading(doc, "3.4.1  Frontend Technologies", 2)

body_para(doc, [
    plain("The frontend of the CVCC system was built using "),
    bold("Next.js 16"),
    plain(" "),
    cite([4]),
    plain(" with the App Router, running on "),
    bold("React 19"),
    plain(" "),
    cite([5]),
    plain(". Next.js was selected because it provides a production-grade React framework with "
          "built-in routing, optimised asset bundling, and strong TypeScript support, all of "
          "which were necessary for a project of this complexity. The App Router's "
          "directory-based routing model made it straightforward to structure the seven role "
          "dashboards as distinct areas of the application under clear paths ("),
    code("/app/hq/dashboard"),
    plain(", "),
    code("/app/facility"),
    plain(", "),
    code("/app/chw"),
    plain(", "),
    code("/app/parent/dashboard"),
    plain(", and so on)."),
], indent=True)

body_para(doc, [
    plain("Styling is handled by "),
    bold("Tailwind CSS v4"),
    plain(" "),
    cite([7]),
    plain(", a utility-first CSS framework that allowed the team to compose component styles "
          "directly in JSX without managing separate stylesheet files. This significantly reduced "
          "the cognitive overhead of maintaining visual consistency. The UI component library "
          "used throughout the application is "),
    bold("shadcn/ui"),
    plain(" "),
    cite([8]),
    plain(", a collection of accessible, headless components built on Radix UI primitives. "
          "Because shadcn/ui components ship as source files rather than a dependency import, "
          "the team was able to customise every component to match the CVCC design language "
          "without fighting against an opinionated component API."),
], indent=True)

body_para(doc, [
    plain("Data visualisation in the dashboards—coverage rates, dose completion trends, and "
          "facility performance metrics—is rendered using "),
    bold("Recharts"),
    plain(" "),
    cite([9]),
    plain(", which integrates naturally with React's component model and handled the dynamic, "
          "server-fetched chart data with minimal configuration."),
], indent=True)

add_heading(doc, "3.4.2  Offline Storage and Mapping", 2)

body_para(doc, [
    plain("The offline storage backbone of the CHW module is "),
    bold("Dexie.js v4"),
    plain(" "),
    cite([10]),
    plain(", a TypeScript-friendly wrapper around the browser's IndexedDB API "),
    cite([19]),
    plain(". Dexie was chosen over raw IndexedDB because it provides a clean, promise-based "
          "API that supports typed schemas, index-based queries, and bulk operations, all of "
          "which were needed to efficiently manage the "),
    code("children"),
    plain(" and "),
    code("vaccinationQueue"),
    plain(" tables while the device is offline."),
], indent=True)

body_para(doc, [
    plain("Geographical mapping for CHW outreach scheduling is powered by "),
    bold("Leaflet"),
    plain(" "),
    cite([11]),
    plain(" and its React wrapper React-Leaflet, using "),
    bold("OpenStreetMap"),
    plain(" "),
    cite([12]),
    plain(" as the tile source. Leaflet was selected over heavier alternatives such as Google "
          "Maps because it is open source, requires no API key, and most importantly, its tile "
          "images are plain cacheable HTTP resources that can be intercepted and stored by the "),
    code("chw-map-sw.js"),
    plain(" service worker. The "),
    code("components/chw/outreach-map.tsx"),
    plain(" component renders the map within CHW outreach sessions, showing the catchment area "
          "and the child's household GPS coordinates. The database also uses Supabase's PostGIS "
          "extension "),
    cite([13]),
    plain(" to store catchment area boundaries as "),
    code("POLYGON"),
    plain(" geometry types, enabling spatial queries on the server side."),
], indent=True)

add_heading(doc, "3.4.3  Backend Framework", 2)

body_para(doc, [
    plain("The server-side application is built with "),
    bold("NestJS 10"),
    plain(" "),
    cite([6]),
    plain(", a progressive Node.js framework built on TypeScript that imposes a structured, "
          "opinionated architecture of modules, controllers, services, and providers. NestJS "
          "was chosen because its architecture closely mirrors the role-based structure needed "
          "by this project: each user role maps naturally to a dedicated NestJS module "
          "("),
    code("AuthModule"),
    plain(", "),
    code("FacilityModule"),
    plain(", "),
    code("ChwModule"),
    plain(", "),
    code("ParentModule"),
    plain(", "),
    code("PhaModule"),
    plain(", "),
    code("BranchManagerModule"),
    plain("), keeping the backend codebase organised and easy to navigate. TypeScript is used "
          "throughout the backend, with strict typing applied to all DTOs (Data Transfer "
          "Objects), service inputs, and database query return types."),
], indent=True)

body_para(doc, [
    plain("The API follows RESTful design conventions with predictable resource-oriented "
          "endpoints. All incoming request bodies are validated using "),
    code("class-validator"),
    plain(" decorators "),
    cite([24]),
    plain(" on DTO classes, ensuring that malformed input is rejected at the boundary of the "
          "application before reaching any business logic. Security headers are applied globally "
          "via the "),
    code("helmet"),
    plain(" middleware, and authentication is implemented using "),
    code("@nestjs/passport"),
    plain(" with a JWT strategy "),
    cite([21]),
    plain("."),
], indent=True)

add_heading(doc, "3.4.4  Database and Background Job Management", 2)

body_para(doc, [
    plain("The persistent data store for the CVCC system is a "),
    bold("PostgreSQL"),
    plain(" database hosted and managed through "),
    bold("Supabase"),
    plain(" "),
    cite([14]),
    plain(". The schema was designed as a fully relational model with explicitly defined "
          "foreign key constraints, enum types, and indexes. The database contains nineteen "
          "core tables, covering user accounts, healthcare facility branches, catchment areas, "
          "guardian records, child records, vaccination events, vaccine schedules, AEFI "
          "(Adverse Event Following Immunisation) reports, certificates, appointments, "
          "notifications, visit logs, offline sync queues, and data deduplication candidates. "
          "The spatial extension "),
    bold("PostGIS"),
    plain(" "),
    cite([13]),
    plain(" was enabled to support the storage of catchment-area polygons as "),
    code("POLYGON"),
    plain(" geometry values and GPS coordinates as "),
    code("POINT"),
    plain(" values, enabling precise geospatial assignment of children and CHWs to their "
          "respective service areas."),
], indent=True)

body_para(doc, [
    plain("Background job scheduling does not use an external queue server. Instead, it is "
          "handled by the "),
    code("@nestjs/schedule"),
    plain(" module, which provides "),
    code("@Cron"),
    plain(" decorator-based scheduling within the NestJS application process. The primary "
          "scheduled job is in "),
    code("VaccinationSchedulerService"),
    plain(" ("),
    code("backend/src/common/vaccination-scheduler.service.ts"),
    plain("), which runs daily at 08:00 West Africa Time (configured for the "),
    code("Africa/Accra"),
    plain(" timezone). On each trigger, it queries the database for all children with "
          "vaccination doses due on that calendar date, retrieves the associated guardian phone "
          "numbers, and dispatches an SMS reminder for each one via the Hubtel SMS API "),
    cite([22]),
    plain(". Every sent notification is recorded in the "),
    code("notifications"),
    plain(" table, creating a durable log of all outbound communications."),
], indent=True)

add_heading(doc, "3.4.5  Utility Libraries", 2)

body_para(doc, [
    plain("Several utility libraries were integrated to support specific functional "
          "requirements. "),
    bold("html5-qrcode v2.3.8"),
    plain(" "),
    cite([16]),
    plain(" underpins the QR code scanning mechanism in the PHA certificate verification "
          "portal. The library's "),
    code("Html5Qrcode"),
    plain(" class is wrapped in a custom "),
    code("QrScanner"),
    plain(" React component ("),
    code("components/pha/qr-scanner.tsx"),
    plain(") that manages camera lifecycle, prefers the device's rear-facing camera, and "
          "handles teardown cleanly on component unmount to avoid holding the camera open "
          "unnecessarily."),
], indent=True)

body_para(doc, [
    plain("Certificate generation relies on "),
    bold("jsPDF v3.0.3"),
    plain(" "),
    cite([15]),
    plain(", a client-side PDF generation library. The "),
    code("generateCertificatePdf()"),
    plain(" function in "),
    code("lib/certificate-pdf.ts"),
    plain(" programmatically constructs an A4-sized vaccination certificate document, drawing "
          "the Ministry of Health branding, the child's full details, a list of all completed "
          "vaccines, the certificate ID (formatted as "),
    code("CERT-GH-YYYY-XXXXXX"),
    plain("), and an embedded QR code image that encodes a URL to the PHA verification "
          "endpoint. The completed PDF is then offered to the user as a browser download. For "
          "QR code generation on the frontend, "),
    bold("qrcode.react v4.2.0"),
    plain(" "),
    cite([17]),
    plain(" is used to render a scannable QR image that is first converted to a data URL "
          "before being passed to jsPDF for embedding."),
], indent=True)

# ─── 3.5 ─────────────────────────────────────────────────────────────────────
add_heading(doc, "3.5  Core System Modules and Implementation", 1)

add_heading(doc, "3.5.1  Identity Verification and QR Code Engine", 2)

body_para(doc, [
    plain("Each child registered in the CVCC system is assigned a unique human-readable "
          "identifier formatted as "),
    code("CH-YYYY-NNN"),
    plain(" (for example, "),
    code("CH-2025-001"),
    plain("), stored in the "),
    code("cvcc_id"),
    plain(" field of the "),
    code("children"),
    plain(" table. In addition to this identifier, each child record carries a "),
    code("qr_code_payload"),
    plain(" field that holds a signed token linking the QR image directly and unambiguously "
          "to that child's record in the database. This QR payload is generated at the time "
          "of child registration and remains immutable unless explicitly reissued."),
], indent=True)

body_para(doc, [
    plain("On the scanning side, the "),
    code("Html5Qrcode"),
    plain(" "),
    cite([16]),
    plain(" component in the PHA portal captures the QR code using the device camera and "
          "extracts the embedded certificate ID or child identifier. This value is then "
          "submitted to the "),
    code("GET /api/pha/certificates/verify"),
    plain(" endpoint on the backend. Before the server processes the query, the certificate "
          "ID parameter is sanitised with a strict regular expression that allows only "
          "alphanumeric characters and hyphens, with a maximum length of one hundred "
          "characters, preventing any injection via the query string. The backend then looks "
          "up the corresponding record and returns the verification result, which the PHA "
          "officer can read on screen."),
], indent=True)

add_heading(doc, "3.5.2  Catchment Area Management and Child Transfer", 2)

body_para(doc, [
    plain("The transfer system allows a child to be moved from the register of one healthcare "
          "facility to another in a controlled and auditable way. Two endpoints in the CHW "
          "controller handle this flow: "),
    code("POST /api/chw/children/:childId/transfer-out"),
    plain(" and "),
    code("POST /api/chw/children/:childId/transfer-in"),
    plain(". The data contracts for these operations are defined in "),
    code("backend/src/chw/dto/transfer.dto.ts"),
    plain(" as "),
    code("TransferOutDto"),
    plain(" and "),
    code("TransferInDto"),
    plain(". When a transfer-out is initiated, the child's "),
    code("primary_facility_id"),
    plain(" in the database is updated to reflect their new destination facility and an audit "
          "record is created. The child's complete vaccination history travels with the record; "
          "nothing is erased or archived."),
], indent=True)

body_para(doc, [
    plain("In the offline CHW context, transfers are queued as "),
    code("VaccinationQueueItem"),
    plain(" entries in the Dexie "),
    cite([10]),
    plain(" "),
    code("vaccinationQueue"),
    plain(" table with action types of "),
    code("transfer_in"),
    plain(" or "),
    code("transfer_out"),
    plain(". The queue item is processed the next time the device has network access and the "),
    code("ChwBackgroundSyncService"),
    plain(" runs. The "),
    code("removeChildFromLocalRegister()"),
    plain(" function in "),
    code("lib/chw-offline/db.ts"),
    plain(" removes the child from the local device register once the transfer-out has been "
          "confirmed by the server, preventing any further offline actions on a record the CHW "
          "no longer holds."),
], indent=True)

add_heading(doc, "3.5.3  Automated Notification System", 2)

body_para(doc, [
    plain("The CVCC system dispatches notifications through two independent channels: SMS for "
          "immediate mobile delivery and email for formal written communication. These two "
          "channels are served by separate services in the NestJS backend. The "),
    code("SmsService"),
    plain(" in "),
    code("backend/src/common/sms.service.ts"),
    plain(" communicates with the "),
    bold("Hubtel SMS API"),
    plain(" "),
    cite([22]),
    plain(", a Ghanaian SMS gateway. It automatically normalises phone numbers into the "
          "international format expected by Hubtel (converting leading-zero numbers to the "),
    code("233xx"),
    plain(" country code format). The service provides four distinct message types: welcome "
          "SMS (sent when a new parent account is created), registration SMS (confirming a "
          "child registration), vaccination reminder (daily outgoing from the cron scheduler), "
          "and appointment confirmation."),
], indent=True)

body_para(doc, [
    plain("The "),
    code("EmailService"),
    plain(" in "),
    code("backend/src/common/email.service.ts"),
    plain(" uses the "),
    bold("Brevo HTTP API"),
    plain(" "),
    cite([23]),
    plain(" ("),
    code("https://api.brevo.com/v3/smtp/email"),
    plain("), called directly via "),
    code("axios"),
    plain(". Email is used for two workflows: a welcome email that delivers a parent's system "
          "credentials upon registration, and a password-reset email that delivers a "
          "time-limited reset link. Both email templates are constructed as full HTML documents "
          "with Ghana Ministry of Health branding embedded directly in the template string."),
], indent=True)

body_para(doc, [
    plain("The daily vaccination reminder cycle is managed by the "),
    code("VaccinationSchedulerService"),
    plain(", which fires at 08:00 Ghana time each day, queries for children with vaccinations "
          "due that day, and calls "),
    code("SmsService.sendVaccinationReminder()"),
    plain(" for each guardian phone number found. A record of every dispatched notification "
          "is written to the "),
    code("notifications"),
    plain(" table, with fields for the channel used, the delivery status, and a reference to "
          "the originating template."),
], indent=True)

add_heading(doc, "3.5.4  Cryptographic Certificate Generation", 2)

body_para(doc, [
    plain("When a child completes all scheduled vaccinations in the CVCC system, the platform "
          "generates a tamper-evident PDF vaccination certificate. This process is implemented "
          "in "),
    code("lib/certificate-pdf.ts"),
    plain(" using the "),
    bold("jsPDF"),
    plain(" library "),
    cite([15]),
    plain(" and is triggered from the parent portal's certificates page. The "),
    code("generateCertificatePdf()"),
    plain(" function compiles a structured A4 document that includes the child's full name, "
          "date of birth, and assigned CVCC identifier; the issuing facility's name; the date "
          "issued; a complete list of vaccines administered and their respective dates; the "
          "certificate's unique ID (formatted as "),
    code("CERT-GH-YYYY-XXXXXX"),
    plain("); and an embedded QR code image generated by qrcode.react "),
    cite([17]),
    plain("."),
], indent=True)

body_para(doc, [
    plain("The embedded QR code is generated from the "),
    code("certificates.qr_payload"),
    plain(" value stored in the database, which encodes a URL to the PHA certificate "
          "verification endpoint. This means that any party—a school admissions officer, a "
          "travel health clinic, or another healthcare provider—can independently verify the "
          "certificate's authenticity by scanning the QR code on the printed document, without "
          "relying on the issuing facility's paper records. The completed PDF is offered as a "
          "direct browser download named using the certificate ID."),
], indent=True)

# ─── 3.6 ─────────────────────────────────────────────────────────────────────
add_heading(doc, "3.6  System Testing, Debugging, and Quality Assurance", 1)

add_heading(doc, "3.6.1  Testing Strategies", 2)

body_para(doc, [
    plain("The quality assurance strategy for the CVCC system was primarily composed of manual "
          "functional testing, endpoint-level integration testing, and structured role-by-role "
          "walkthroughs. Each time a new feature was added during an iteration, the responsible "
          "developer performed a full walkthrough of the feature from the frontend UI down to "
          "the database, verifying that data entered in the browser appeared correctly in the "
          "Supabase dashboard "),
    cite([14]),
    plain(" and that API responses matched the expected schemas."),
], indent=True)

body_para(doc, [
    plain("For the backend, integration testing of the NestJS API endpoints was performed "
          "using HTTP client tools that allowed the team to send crafted requests to each "
          "endpoint with different payloads, authentication tokens, and edge-case inputs. This "
          "was particularly important for endpoints that involve multiple database tables, such "
          "as the child registration flow, which creates or links records in the "),
    code("users"),
    plain(", "),
    code("guardians"),
    plain(", "),
    code("children"),
    plain(", and "),
    code("child_guardian"),
    plain(" tables in a single operation. DTO-level validation "),
    cite([24]),
    plain(" was also verified by deliberately sending malformed request bodies and confirming "
          "that "),
    code("class-validator"),
    plain(" responded with the correct 400 Bad Request errors before any database write was "
          "attempted."),
], indent=True)

body_para(doc, [
    plain("The NestJS test infrastructure ("),
    code("@nestjs/testing"),
    plain(", "),
    code("jest"),
    plain(", "),
    code("ts-jest"),
    plain(") is installed as a development dependency, providing the scaffolding needed to "
          "add unit and end-to-end tests in future iterations of the project."),
], indent=True)

add_heading(doc, "3.6.2  Offline Synchronisation Testing", 2)

body_para(doc, [
    plain("Testing the CHW offline module required a different approach from the rest of the "
          "system, because the failure mode being tested—losing network connectivity "
          "mid-session—cannot be replicated by simply submitting a bad form value. The team "
          "used the browser's DevTools Network panel to simulate offline conditions by setting "
          "the network throttling preset to \"Offline\" while the CHW dashboard was active."),
], indent=True)

body_para(doc, [
    plain("With the device simulated as offline, test scenarios were executed: registering a "
          "new child, recording a vaccination event, and initiating a child transfer. After "
          "each action, the contents of the "),
    code("cvcc_chw_offline_v2"),
    plain(" IndexedDB "),
    cite([19]),
    plain(" were inspected directly in the browser's Application panel to confirm that the "
          "records were correctly written and that the "),
    code("vaccinationQueue"),
    plain(" table contained the corresponding pending items with the correct "),
    code("idempotencyKey"),
    plain(" values. The device was then returned to \"Online\" status to simulate network "
          "restoration. The "),
    code("ChwBackgroundSyncService"),
    plain(" was observed to drain the queue within its five-minute polling window, and the "
          "Supabase database was checked to confirm that all queued records had been reconciled "
          "accurately on the server. Edge cases tested included: power-cycling the browser "
          "before sync (verifying that IndexedDB records persist across sessions), deliberately "
          "repeating a sync call (verifying that idempotency keys prevented duplicate "
          "insertions), and simulating a failed sync for individual items (verifying that "),
    code("retryCount"),
    plain(" and "),
    code("lastError"),
    plain(" were updated correctly in the queue)."),
], indent=True)

add_heading(doc, "3.6.3  Bug Tracking and Resolution", 2)

body_para(doc, [
    plain("Bugs identified during the development process were tracked as issues in the "
          "project's version control repository and documented in the "),
    code("CHANGELOG.md"),
    plain(" file. This provided a chronological record of what was discovered, what the root "
          "cause was, and what change resolved it."),
], indent=True)

body_para(doc, [
    plain("Several notable bugs were encountered and resolved during the iterative cycles. "
          "The password reset flow initially returned a generic authentication failure when a "
          "user submitted a valid reset token that had expired, because the expiry check was "
          "performed on the database timestamp without accounting for timezone differences "
          "between the server process and the Supabase database. This was resolved by "
          "normalising all timestamp comparisons to UTC within the "),
    code("AuthService"),
    plain(". The email service was originally configured to use an SMTP-based transporter, "
          "but integration testing revealed that the SMTP credentials for the Brevo service "
          "were not accepted reliably in the deployment environment; the service was "
          "subsequently rewritten to use Brevo's HTTP API "),
    cite([23]),
    plain(" directly via "),
    code("axios"),
    plain(", which proved more stable. The "),
    code("reset-password"),
    plain(" page also triggered a React hydration error in production builds because it called "
          "the "),
    code("useSearchParams"),
    plain(" hook from the Next.js "),
    cite([4]),
    plain(" App Router at the top level of the component tree; this was resolved by wrapping "
          "the component in a "),
    code("Suspense"),
    plain(" boundary, as logged in the commit history."),
], indent=True)

# ─── 3.7 CONCLUSION ──────────────────────────────────────────────────────────
add_heading(doc, "3.7  Conclusion", 1)

body_para(doc, [
    plain("This chapter has documented the methodology and design decisions that underpin the "
          "CVCC system. The semi-structured interview process "),
    cite([2]),
    plain(" surfaced the field constraints—network unavailability, catchment area rules, and "
          "cross-facility transfers—that directly shaped the architecture. The Agile methodology "
          ""),
    cite([1]),
    plain(" allowed the three-person team to deliver functional role slices incrementally and "
          "to incorporate discoveries from each iteration into the next. The three-tier "
          "architecture (Next.js "),
    cite([4]),
    plain(" PWA, NestJS "),
    cite([6]),
    plain(" API, Supabase PostgreSQL "),
    cite([14]),
    plain(") provides a clean and maintainable separation of concerns, while the offline-first "
          "CHW module—built on Dexie.js "),
    cite([10]),
    plain(", two purpose-built service workers, and AES-GCM-256 field encryption "),
    cite([20]),
    plain("—ensures that data collection continues reliably even in the most resource-constrained "
          "environments. The role-based access control system, enforced at the JWT "),
    cite([21]),
    plain(", API guard, and RLS levels simultaneously, guarantees that the seven user roles "
          "remain securely isolated from one another. Together, the technology stack and "
          "architectural choices described here provided a solid and well-evidenced foundation "
          "for the implementation phase described in Chapter Four."),
], indent=True)

# ─── REFERENCES ──────────────────────────────────────────────────────────────
doc.add_page_break()
add_heading(doc, "References", 1)

references = [
    ('[1]', 'K. Beck, M. Beedle, A. van Bennekum, A. Cockburn, W. Cunningham, M. Fowler, '
            'J. Grenning, J. Highsmith, A. Hunt, R. Jeffries, J. Kern, B. Marick, R. C. '
            'Martin, S. Mellor, K. Schwaber, J. Sutherland, and D. Thomas, "Manifesto for '
            'Agile Software Development," 2001. [Online]. Available: '
            'https://agilemanifesto.org. [Accessed: Mar. 15, 2026].'),
    ('[2]', 'I. Seidman, Interviewing as Qualitative Research: A Guide for Researchers in '
            'Education and the Social Sciences, 4th ed. New York, NY, USA: Teachers College '
            'Press, 2013.'),
    ('[3]', 'World Health Organization, "Immunization Agenda 2030: A Global Strategy to '
            'Leave No One Behind," WHO, Geneva, Switzerland, Tech. Rep., 2020. [Online]. '
            'Available: https://www.who.int/initiatives/immunization-agenda-2030. '
            '[Accessed: Mar. 15, 2026].'),
    ('[4]', 'Vercel, "Next.js Documentation," 2024. [Online]. Available: '
            'https://nextjs.org/docs. [Accessed: Mar. 15, 2026].'),
    ('[5]', 'Meta Open Source, "React – A JavaScript Library for Building User Interfaces," '
            '2024. [Online]. Available: https://react.dev. [Accessed: Mar. 15, 2026].'),
    ('[6]', 'NestJS Core Team, "NestJS – A Progressive Node.js Framework," 2024. [Online]. '
            'Available: https://nestjs.com. [Accessed: Mar. 15, 2026].'),
    ('[7]', 'A. Wathan and T. Kennedy, "Tailwind CSS – A Utility-First CSS Framework," '
            '2024. [Online]. Available: https://tailwindcss.com. [Accessed: Mar. 15, 2026].'),
    ('[8]', 'shadcn, "shadcn/ui – Beautifully Designed Components Built with Radix UI and '
            'Tailwind CSS," 2024. [Online]. Available: https://ui.shadcn.com. '
            '[Accessed: Mar. 15, 2026].'),
    ('[9]', 'Recharts Group, "Recharts – A Redefined Chart Library Built with React and '
            'D3," 2024. [Online]. Available: https://recharts.org. '
            '[Accessed: Mar. 15, 2026].'),
    ('[10]', 'D. Fahlander, "Dexie.js – A Minimalistic Wrapper for IndexedDB," 2024. '
             '[Online]. Available: https://dexie.org. [Accessed: Mar. 15, 2026].'),
    ('[11]', 'V. Agafonkin, "Leaflet – An Open-Source JavaScript Library for '
             'Mobile-Friendly Interactive Maps," 2024. [Online]. Available: '
             'https://leafletjs.com. [Accessed: Mar. 15, 2026].'),
    ('[12]', 'OpenStreetMap Foundation, "OpenStreetMap," 2024. [Online]. Available: '
             'https://www.openstreetmap.org. [Accessed: Mar. 15, 2026].'),
    ('[13]', 'PostGIS Development Team, "PostGIS – Spatial and Geographic Objects for '
             'PostgreSQL," 2024. [Online]. Available: https://postgis.net. '
             '[Accessed: Mar. 15, 2026].'),
    ('[14]', 'Supabase Inc., "Supabase – The Open Source Firebase Alternative," 2024. '
             '[Online]. Available: https://supabase.com. [Accessed: Mar. 15, 2026].'),
    ('[15]', 'MrRio, "jsPDF – Client-Side JavaScript PDF Generation for the Browser," '
             '2024. [Online]. Available: https://github.com/parallax/jsPDF. '
             '[Accessed: Mar. 15, 2026].'),
    ('[16]', 'M. Dhar, "html5-qrcode – A Lightweight QR Code and Bar Code Scanning '
             'Library for the Web," 2024. [Online]. Available: '
             'https://github.com/mebjas/html5-qrcode. [Accessed: Mar. 15, 2026].'),
    ('[17]', 'U. Khan, "qrcode.react – A QR Code Component for React Applications," '
             '2024. [Online]. Available: https://github.com/zpao/qrcode.react. '
             '[Accessed: Mar. 15, 2026].'),
    ('[18]', 'Google Developers, "Progressive Web Apps," Web.dev, 2023. [Online]. '
             'Available: https://web.dev/progressive-web-apps. [Accessed: Mar. 15, 2026].'),
    ('[19]', 'W3C, "Indexed Database API 3.0," W3C Working Draft, 2022. [Online]. '
             'Available: https://www.w3.org/TR/IndexedDB. [Accessed: Mar. 15, 2026].'),
    ('[20]', 'National Institute of Standards and Technology (NIST), "Advanced Encryption '
             'Standard (AES)," Federal Information Processing Standards Publication '
             'FIPS PUB 197, Nov. 2001.'),
    ('[21]', 'M. Jones, J. Bradley, and N. Sakimura, "JSON Web Token (JWT)," '
             'IETF RFC 7519, May 2015. [Online]. Available: '
             'https://tools.ietf.org/html/rfc7519. [Accessed: Mar. 15, 2026].'),
    ('[22]', 'Hubtel, "Hubtel SMS API Documentation," 2024. [Online]. Available: '
             'https://developers.hubtel.com. [Accessed: Mar. 15, 2026].'),
    ('[23]', 'Brevo, "Brevo Transactional Email API," 2024. [Online]. Available: '
             'https://developers.brevo.com. [Accessed: Mar. 15, 2026].'),
    ('[24]', 'TypeStack, "class-validator – Decorator-Based Property Validation for '
             'TypeScript/JavaScript," 2024. [Online]. Available: '
             'https://github.com/typestack/class-validator. [Accessed: Mar. 15, 2026].'),
    ('[25]', 'Web Platform Incubator Community Group (WICG), "Web Background '
             'Synchronization Specification," W3C, 2021. [Online]. Available: '
             'https://wicg.github.io/background-sync/spec. [Accessed: Mar. 15, 2026].'),
]

for num, text in references:
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=2, after=4)
    pf = p.paragraph_format
    pf.left_indent        = Inches(0.4)
    pf.first_line_indent  = Inches(-0.4)   # hanging indent

    r_num = p.add_run(num + "  ")
    r_num.font.name = "Times New Roman"
    r_num.font.size = Pt(11)
    r_num.font.bold = True

    r_text = p.add_run(text)
    r_text.font.name = "Times New Roman"
    r_text.font.size = Pt(11)

# ─── SAVE ─────────────────────────────────────────────────────────────────────
out_path = r"c:\Users\Junior Owusu\Desktop\Final Year Project\child-vaccination-system\docs\CHAPTER-3-IEEE-FULL.docx"
doc.save(out_path)
print("Saved: " + out_path)

